"""
Generate APA 7th Style Word Document for IJDL Design Case Paper
Based on FINAL_outline_v2.md
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BASE = "/Volumes/External SSD/Projects/Diverga/diverga-evaluation-paper/Paper1_DesignCase"
FIG_DIR = os.path.join(BASE, "figures")
OUTPUT = os.path.join(BASE, "You_Yang_2026_Design_Case_FINAL.docx")

doc = Document()

# --- Page setup: APA 7th (1-inch margins, Letter) ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 2.0

# --- Helper functions ---
def add_heading_apa(doc, text, level=1):
    """APA 7th heading styles"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if level == 1:
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.bold = True
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_figure(doc, path, caption, fig_num):
    """Insert figure with APA caption"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    if os.path.exists(path):
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
    else:
        run = p.add_run(f"[IMAGE PLACEHOLDER: {os.path.basename(path)}]")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    # Caption
    cap = doc.add_paragraph()
    cap.paragraph_format.line_spacing = 2.0
    cap.paragraph_format.space_after = Pt(6)
    r1 = cap.add_run(f"Figure {fig_num}")
    r1.bold = True
    r1.italic = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    cap.add_run("\n")
    r2 = cap.add_run(caption)
    r2.italic = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return cap

def add_apa_table(doc, headers, rows, table_num, title):
    """Create APA 7th style table"""
    # Table number and title
    p_num = doc.add_paragraph()
    p_num.paragraph_format.line_spacing = 2.0
    p_num.paragraph_format.space_before = Pt(12)
    r = p_num.add_run(f"Table {table_num}")
    r.bold = True
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.line_spacing = 2.0
    p_title.paragraph_format.space_after = Pt(6)
    r = p_title.add_run(title)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    # Remove vertical borders, keep horizontal (APA style)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'bottom', 'insideH']:
        b = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '000000'
        })
        borders.append(b)
    for border_name in ['left', 'right', 'insideV']:
        b = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none',
            qn('w:sz'): '0',
            qn('w:space'): '0',
            qn('w:color'): '000000'
        })
        borders.append(b)
    tblPr.append(borders)

    doc.add_paragraph()  # spacing after table
    return table


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph().paragraph_format.line_spacing = 2.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("Scaffolding Without Replacing: A Design Case of Human Checkpoints\nand Verbalized Sampling in an AI Research Methodology Assistant")
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

doc.add_paragraph().paragraph_format.line_spacing = 2.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("Hosung You")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("College of Education, Pennsylvania State University")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

doc.add_paragraph().paragraph_format.line_spacing = 2.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("Mohan Yang")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("College of Education and Human Development, Texas A&M University")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

doc.add_paragraph().paragraph_format.line_spacing = 2.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("Author Note")
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

add_body(doc, "Hosung You is a doctoral student in the College of Education at Pennsylvania State University. Mohan Yang is an assistant professor in the College of Education and Human Development at Texas A&M University. Correspondence concerning this article should be addressed to Hosung You, College of Education, Pennsylvania State University, University Park, PA 16802.", indent=False)

# ============================================================
# PAGE BREAK -> ABSTRACT
# ============================================================
doc.add_page_break()

add_heading_apa(doc, "Abstract", level=1)

add_body(doc, (
    "AI research tools increasingly shape doctoral researchers' methodological decisions, yet these tools exhibit mode collapse, "
    "converging on predictable recommendations regardless of research context, while encouraging automation bias through uncritical acceptance. "
    "This first-person design case documents Diverga, a 24-agent AI research methodology assistant designed to scaffold rather than replace "
    "researchers' methodological reasoning. Drawing on cognitive apprenticeship (Collins, Brown, & Newman, 1989), we identify how Diverga's "
    "architecture operationalizes six learning methods (modeling, coaching, scaffolding, articulation, reflection, and exploration) through "
    "AI-mediated mechanisms rather than explicit instruction. Three design decisions are examined: (a) a multi-agent architecture consolidated "
    "from 44 to 24 agents through principled scaffolding calibration, (b) Verbalized Sampling with Typicality Scores that structurally enforce "
    "recommendation diversity and prompt metacognitive reflection, and (c) five required Human Checkpoints functioning as cognitive forcing "
    "functions (Bucinca et al., 2021) that compel articulation of decision rationale. The 57-day development process (224 commits) is organized "
    "into five iterative design cycles following the Educational Design Research model (McKenney & Reeves, 2012), revealing how the co-evolution "
    "of platform capabilities and design ambitions shaped the system. Design failures, including a same-day protocol reversal and a silent "
    "checkpoint bug, are documented as evidence of genuine iterative reasoning. Diverga is framed as a learning design in disguise: a research "
    "tool whose architectural form scaffolds the cognitive processes that characterize effective methodological reasoning in the academic workplace."
), indent=False)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 2.0
p.paragraph_format.first_line_indent = Cm(1.27)
r = p.add_run("Keywords: ")
r.italic = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r2 = p.add_run("design case, AI scaffolding, cognitive apprenticeship, mode collapse, human checkpoints, multi-agent systems, research methodology, educational design research")
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)

# ============================================================
# PAGE BREAK -> BODY
# ============================================================
doc.add_page_break()

# Re-state title at start of body (APA 7th)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("Scaffolding Without Replacing: A Design Case of Human Checkpoints\nand Verbalized Sampling in an AI Research Methodology Assistant")
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

# ============================================================
# SECTION 1: INTRODUCTION
# ============================================================
add_heading_apa(doc, "Introduction", level=1)

add_body(doc, (
    "Research methodology decisions are consequential knowledge work tasks for doctoral students and early-career scholars. "
    "Selecting a theoretical framework, choosing between qualitative and quantitative approaches, or determining an appropriate "
    "sampling strategy are not clerical operations; they define the epistemic character of a study and shape a researcher's "
    "professional trajectory. When a doctoral student asks an AI tool for a theoretical framework recommendation and receives "
    "the Technology Acceptance Model regardless of context, the tool has not assisted methodological reasoning; it has replaced it "
    "with a statistical default."
), indent=True)

add_body(doc, (
    "This pattern, which we term mode collapse following Zhao et al. (2024), describes the tendency of AI systems trained to "
    "maximize prediction probability to converge on the most statistically common output rather than the most contextually "
    "appropriate one. The phenomenon is not a bug in any particular tool; it is a structural property of how large language models "
    "are trained. The most probable response is, by definition, the most common one. For researchers seeking methodological "
    "creativity and contextual fit, the most common recommendation is often not the best one."
))

add_body(doc, (
    "Mode collapse is compounded by automation bias: the tendency to accept AI recommendations without critical evaluation, "
    "especially under time pressure. A 2025 Springer Nature report found a 78% uncritical acceptance rate for AI-generated "
    "methodology suggestions, and Microsoft Research (2025) identified a negative correlation between AI tool use and critical "
    "thinking engagement in knowledge work settings. Together, mode collapse narrows the options presented, and automation bias "
    "ensures the narrowed options are accepted without scrutiny."
))

add_body(doc, (
    "The design challenge, then, is this: Can an AI research assistant be designed to scaffold methodological reasoning rather "
    "than replace it? This is a learning design problem disguised as an engineering problem."
))

add_body(doc, (
    "Diverga is, at one level, an AI research tool. At another level, it is a learning design embedded in a tool's architecture. "
    "To make this claim precise, we draw on the cognitive apprenticeship framework (Collins, Brown, & Newman, 1989), which "
    "identifies six methods by which learning environments make expert thinking visible to developing practitioners: modeling, "
    "coaching, scaffolding, articulation, reflection, and exploration. In traditional cognitive apprenticeship, a human expert "
    "models reasoning, coaches the learner, and gradually fades support. Diverga implements these methods not through human "
    "mentorship but through structural design features."
))

add_body(doc, (
    "The multi-agent architecture models expert reasoning by distributing research design across specialized agents, each "
    "embodying domain-specific methodological knowledge. The Verbalized Sampling system, with its Typicality Scores, supports "
    "reflection and exploration by making the statistical commonality of each recommendation transparent. The Human Checkpoint "
    "system supports articulation by requiring the researcher to state and justify decisions at five mandatory interruption "
    "points, which function as cognitive forcing functions (Bucinca et al., 2021) that shift processing from automatic acceptance "
    "to deliberate evaluation. In the special section's terms, Diverga is a learning design for the workplace of doctoral research: "
    "it shapes the performance of methodological judgment through structural scaffolding rather than didactic transmission."
))

add_body(doc, (
    "This paper documents the design process, decisions, and rationale across five iterative design cycles over 57 days (224 "
    "version-controlled commits). Empirical evaluation of Diverga's effects on researcher performance is reserved for a companion "
    "experimental paper. The design case follows Educational Design Research phases 1 and 2 (McKenney & Reeves, 2012), and the "
    "first-person narrative follows IJDL design case conventions (Boling, 2010; Smith, 2010). Table 3 maps all six cognitive "
    "apprenticeship methods to specific Diverga mechanisms."
))

# Table 3: CA Mapping
add_apa_table(doc,
    headers=["CA Method", "Diverga Mechanism", "Design Decision"],
    rows=[
        ["Modeling", "Agent-generated expert reasoning; VS Arena multi-perspective demonstration", "Multi-Agent Architecture (5.1)"],
        ["Coaching", "Agent-specific contextual feedback; domain-calibrated guidance", "Multi-Agent Architecture (5.1)"],
        ["Scaffolding", "9-category task decomposition; fading via 44-to-24 consolidation", "Multi-Agent Architecture (5.1)"],
        ["Articulation", "Checkpoint-required decision rationale; Decision Audit Trail", "Human Checkpoints (5.3)"],
        ["Reflection", "T-Score metacognitive prompts; typicality transparency", "VS and T-Scores (5.2)"],
        ["Exploration", "Long-tail sampling (T < 0.4); VS Arena epistemological diversity", "VS and T-Scores (5.2)"],
    ],
    table_num=3,
    title="Cognitive Apprenticeship Methods Mapped to Diverga Design Mechanisms"
)

# ============================================================
# SECTION 2: DESIGN CONTEXT
# ============================================================
add_heading_apa(doc, "Design Context", level=1)

add_heading_apa(doc, "The Designer and the Dual Role", level=2)
add_body(doc, (
    "The first author, a doctoral student in education at Pennsylvania State University, designed and built Diverga over 57 days "
    "(January 22 to March 19, 2026). This dual role as designer and researcher is both an asset, providing insider tacit knowledge "
    "about research design challenges, and a limitation, introducing evaluation bias. Dr. Mohan Yang, an instructional design "
    "scholar at Texas A&M University with published IJDL design cases (Yang & Harbor, 2023), was invited as a collaborating "
    "researcher and critical external voice. This collaboration structure is itself a design decision, documented here rather than "
    "buried in an acknowledgment."
))

add_heading_apa(doc, "Target Users and Workplace Context", level=2)
add_body(doc, (
    "Diverga's target users are social science doctoral students engaged in research design. Research methodology decisions "
    "constitute their primary knowledge work: selecting theoretical frameworks, designing data collection procedures, choosing "
    "analytical approaches. A persistent gap exists between research methods coursework, which teaches the logic of methods, and "
    "actual research design practice, which demands the judgment required to select and adapt methods for a specific study. "
    "Diverga is designed for that gap."
))

add_heading_apa(doc, "Technical Context", level=2)
add_body(doc, (
    "Diverga is built on the Model Context Protocol (MCP), an open standard for connecting AI assistants to external tools and "
    "data sources (Anthropic, 2024). MCP abstracts the complexity of agent-to-agent communication, enabling a doctoral student "
    "to build a multi-agent system in 57 days. Claude Code provides the execution environment, and the plugin ecosystem enables "
    "marketplace distribution. Table 1 compares Diverga with existing AI research tools across five design-relevant dimensions."
))

# Table 1: Comparison
add_apa_table(doc,
    headers=["Dimension", "Diverga", "Elicit", "Consensus", "ChatGPT"],
    rows=[
        ["Architecture", "24 specialized agents in 9 categories", "Single retrieval pipeline", "Single retrieval pipeline", "Single general-purpose agent"],
        ["Diversity mechanism", "Verbalized Sampling with T-Scores", "None (retrieval-based)", "None (retrieval-based)", "Prompt-dependent"],
        ["Human oversight", "5 REQUIRED checkpoint gates", "None (user-initiated)", "None (user-initiated)", "None (user-initiated)"],
        ["Typicality transparency", "T-Score displayed for every recommendation", "N/A", "N/A", "N/A"],
        ["Decision audit trail", "SQLite-logged decisions at each checkpoint", "N/A", "N/A", "Conversation history only"],
    ],
    table_num=1,
    title="Comparison of Diverga with Existing AI Research Tools"
)

# ============================================================
# SECTION 3: DESIGN PROCESS
# ============================================================
add_heading_apa(doc, "Design Process: Five Iterative Cycles", level=1)

add_body(doc, (
    "The 57-day development period is organized into five iterative cycles following the Educational Design Research framework "
    "(McKenney & Reeves, 2012), with each cycle comprising analysis/exploration, design/construction, and evaluation/reflection "
    "phases. These cycles were not pre-planned; they emerged from the interaction between design ambitions, evaluation feedback, "
    "and the evolving capabilities of the Claude Code platform (see Figure 1 for the conceptual framework and Figure 2 for the "
    "five cycles mapped to the EDR model)."
))

# Figure 1: Conceptual Framework
add_figure(doc,
    os.path.join(FIG_DIR, "figure1_conceptual_framework.png"),
    "Conceptual Framework: Educational Design Research Model, Cognitive Apprenticeship Methods, and Three Design Decisions, With Platform Evolution as Environmental Context",
    1
)

add_heading_apa(doc, "Cycle 1: Foundation (v1.0-v2.0; January 22-23, 2026)", level=2)
add_body(doc, (
    "The project began with the recognition that LLM-based research assistants default to modal recommendations. The initial "
    "commit established 20 agents in 5 categories, and Verbalized Sampling methodology (Zhao et al., 2024) was integrated within "
    "the first day, introducing the T-Score system for measuring recommendation typicality. Evaluation revealed marketplace "
    "compatibility issues and overly complex installation, prompting simplification."
))

add_heading_apa(doc, "Cycle 2: Sisyphus and Clean Slate (v3.0-v6.0; January 23-25, 2026)", level=2)
add_body(doc, (
    "Five modular creativity mechanisms were designed. The system expanded to 33 agents. Critically, the Sisyphus Protocol was "
    "introduced in v5.0: an autonomous continuation mechanism that removed human control between checkpoints. This was the "
    "fastest design failure in the project, added and removed on the same day. The v6.0 Clean Slate Edition established the "
    "foundational principle that guided all subsequent development: AI works between checkpoints; humans decide at them. The "
    "project was renamed twice before settling on Diverga."
))

add_heading_apa(doc, "Cycle 3: Expansion to Peak (v6.0-v6.7; January 25-February 4, 2026)", level=2)
add_body(doc, (
    "Domain-specific gaps were identified (meta-analysis, systematic review, humanization), driving expansion to 44 agents. "
    "Cross-platform support was added for Codex CLI and OpenCode. Plugin discovery proved fragile, requiring eight consecutive "
    "fix commits in a single day. Cross-platform parity was never fully achieved (75% for Codex CLI, 70% for OpenCode), "
    "foreshadowing the eventual abandonment of multi-platform support."
))

add_heading_apa(doc, "Cycle 4: The Great Consolidation (v8.0-v11.0; February 4-March 10, 2026)", level=2)
add_body(doc, (
    "With 44 agents across three platforms and a 59KB configuration file, the system had become unwieldy. Infrastructure "
    "matured through MCP checkpoint enforcement, SQLite backend adoption, and Agent Teams pilot integration. Then came the "
    "consolidation: 44 to 24 agents, cross-platform support removed entirely, configuration reduced by 79%. A silent checkpoint "
    "bug was discovered in v11.0 in which all checkpoints auto-approved without user input; the system appeared to have "
    "oversight but did not. The Claude Code hooks system (introduced June 2025) was critical to this cycle: it enabled "
    "architectural enforcement of checkpoints, replacing the behavioral instructions that had proven unreliable."
))

add_heading_apa(doc, "Cycle 5: Stabilization (v11.1-v12.0; March 12-19, 2026)", level=2)
add_body(doc, (
    "The simplified foundation enabled new capabilities. VS Arena was implemented with five epistemological persona agents. "
    "The Claude Code Agent Teams feature (February 2026) was consequential: before Agent Teams, VS Arena required simulated "
    "diversity through subagent calls in the same context window; after Agent Teams, each persona operated in its own "
    "independent context, producing substantively different debate dynamics. The configuration was further slimmed from 16.8KB "
    "to less than 5KB, and redundant orchestration skills were unified."
))

# Figure 2: EDR Cycles
add_figure(doc,
    os.path.join(FIG_DIR, "figure2_edr_cycles.png"),
    "Five Iterative Design Cycles Mapped to the Educational Design Research Model (McKenney & Reeves, 2012), With Agent Count Trajectory and Key Events",
    2
)

add_body(doc, (
    "Figure 3 shows the co-evolution of Claude Code platform features and Diverga design changes. Three platform capabilities "
    "were particularly consequential: hooks (enabling checkpoint enforcement at the architectural level), plugins (enabling "
    "marketplace distribution), and Agent Teams (enabling genuine multi-agent epistemological debate)."
))

# Figure 3: Platform Co-evolution
add_figure(doc,
    os.path.join(FIG_DIR, "figure3_platform_coevolution.png"),
    "Platform-Design Co-evolution: Claude Code Feature Milestones and Corresponding Diverga Design Changes",
    3
)

# ============================================================
# SECTION 4: DESIGN PROBLEM
# ============================================================
add_heading_apa(doc, "Design Problem", level=1)

add_heading_apa(doc, "Mode Collapse in Research AI", level=2)
add_body(doc, (
    "Mode collapse describes the tendency of AI systems to converge on the most statistically common output regardless of "
    "context. When researchers query AI tools for methodology guidance, the tools converge on the same high-probability "
    "recommendations: the Technology Acceptance Model for technology adoption studies, thematic analysis for qualitative research, "
    "random effects models for meta-analysis. This convergence is not random error; it is a systematic artifact of how large "
    "language models are trained to maximize predictive probability (Zhao et al., 2024). An analogy from instructional design "
    "may clarify the concern: mode collapse is to AI methodology guidance what a standardized, one-size-fits-all curriculum is "
    "to pedagogy. Both reduce diversity in the service of efficiency."
))

add_heading_apa(doc, "Automation Bias in Knowledge Work", level=2)
add_body(doc, (
    "Automation bias, the tendency to accept AI recommendations without critical evaluation, compounds mode collapse. A 2025 "
    "Springer Nature report found a 78% uncritical acceptance rate for AI-generated methodology suggestions. Microsoft Research "
    "(2025) identified a negative correlation between AI tool use and critical thinking engagement in knowledge work settings. "
    "The combination narrows the options presented and ensures the narrowed options go unscrutinized."
))

add_heading_apa(doc, "The Design Gap", level=2)
add_body(doc, (
    "No existing tool combines structured human oversight with diversity-generating mechanisms for research methodology guidance. "
    "The design challenge is not to build a better recommendation engine; it is to build a system that keeps the researcher's "
    "judgment in the loop while countering the homogenization pressure of the underlying model. The cognitive apprenticeship "
    "framework clarifies the goal: the system should make expert methodological thinking visible and require the researcher to "
    "engage with it, not merely consume it."
))

# ============================================================
# SECTION 5: DESIGN DECISIONS
# ============================================================
add_heading_apa(doc, "Design Decisions and Rationale", level=1)

# 5.1
add_heading_apa(doc, "Multi-Agent Architecture: From 44 to 24 Agents", level=2)
add_body(doc, (
    "In cognitive apprenticeship terms, the multi-agent architecture operationalizes three methods. Each agent models expert "
    "reasoning in its domain. Domain-specific feedback constitutes coaching. The 9-category task decomposition provides "
    "scaffolding. The consolidation from 44 to 24 is scaffolding calibration: reducing surface complexity so intellectual "
    "structure becomes visible."
))

add_body(doc, (
    "Diverga grew from 20 agents at launch (v1.0) to a peak of 44 agents (v6.7) before being consolidated to 24 agents "
    "across 9 functional categories (v11.0). The trajectory to 44 included rapid expansion driven by domain-specific gaps "
    "(meta-analysis, systematic review, humanization), cross-platform support that multiplied complexity, and a 104-file Python "
    "memory module built over 11 days and then replaced entirely by MCP-based architecture. These expansions were not wasted; "
    "they revealed the thresholds at which complexity becomes counterproductive."
))

add_body(doc, (
    "Three alternatives were considered. A single-agent system with diverse prompting would be simpler but cannot hold competing "
    "epistemological commitments simultaneously. A modular tool-use architecture similar to Cline or Obra maximizes adoption "
    "breadth but sacrifices domain depth. Maintaining 44 agents offers maximum specialization, but coordination overhead and "
    "user cognitive load scale superlinearly with agent count (Li et al., 2025)."
))

add_body(doc, (
    "The consolidation was driven by functional overlap analysis (several v6.7 agents performed structurally identical tasks in "
    "different topical clothing), usage patterns (many specialized agents were rarely invoked), and the principle that "
    "specialization should map to genuine epistemic distinctions rather than surface-level topic divisions. The 9-category "
    "structure (Research Foundation, Literature, Design, Data, Analysis, Quality, Communication, Systematic Review, and Ethics) "
    "maps to the actual phases of research design, not to the internal logic of the AI system."
))

# Figure 4: Architecture Comparison
add_figure(doc,
    os.path.join(FIG_DIR, "figure4_architecture_comparison.png"),
    "Architecture Comparison: v10 (44 Agents) Versus v11 (24 Agents) Across Nine Functional Categories",
    4
)

# Table 2: Agent Consolidation
add_apa_table(doc,
    headers=["Removed Agent(s)", "Merged Into", "Rationale"],
    rows=[
        ["A3 (Devil's Advocate) + A6 (Visualizer)", "A2", "Overlapping critical analysis functions"],
        ["A4 (Ethics) + F4 (Bias Detection)", "X1 (Research Guardian)", "Both serve research integrity"],
        ["B3 (Effect Size) + C6 + C7", "C5", "All serve meta-analysis validation"],
        ["B4 (Research Radar)", "Removed entirely", "Low usage"],
        ["C4 (Materials) + D1 (Sampling)", "C1", "Design-adjacent functions"],
        ["F1 + F2 + F3 (Quality)", "G2", "Quality checks into publication"],
        ["G3 (Peer Review) + G4 (Pre-reg)", "G2", "Publication workflow consolidated"],
        ["H1 (Ethnography) + H2 (Action Research)", "C2", "Qualitative methods consolidated"],
    ],
    table_num=2,
    title="Agent Consolidation Summary: Version 10 to Version 11"
)

# 5.2
add_heading_apa(doc, "Verbalized Sampling and Typicality Scores: Structuring Diversity", level=2)
add_body(doc, (
    "In cognitive apprenticeship terms, the T-Score is a reflection device that prompts metacognitive evaluation: Am I choosing "
    "this because it is best for my study, or because it is common? The VS methodology's long-tail sampling supports exploration "
    "by presenting defensible alternatives the researcher would not encounter through modal convergence alone."
))

add_body(doc, (
    "Asking a curriculum to \"be diverse\" does not produce diversity; structural requirements for diverse content do. Similarly, "
    "prompting an AI to \"provide diverse recommendations\" is an instruction it can partially satisfy while still defaulting to "
    "high-probability outputs. Diversity must be structurally enforced, not instructionally requested."
))

add_body(doc, (
    "Diverga operationalizes Verbalized Sampling (Zhao et al., 2024), a methodology that generates recommendations across the "
    "full typicality spectrum. The Typicality Score (T-Score) measures how common a recommendation is relative to the model's "
    "training distribution: T > 0.8 is modal; T < 0.4 is genuinely unconventional. Critically, the T-Score is surfaced to the "
    "researcher rather than used only internally. A researcher who sees T = 0.82 knows they are receiving the expected "
    "recommendation; one who sees T = 0.31 knows they are in unconventional territory. The T-Score thereby becomes a "
    "metacognitive prompt rather than merely a quality indicator. T-Score calibration uses OpenAlex frequency lookup for "
    "empirical grounding."
))

add_body(doc, (
    "The VS Arena extends this principle through five epistemological persona agents (post-positivist, interpretivist, critical "
    "theorist, pragmatist, transformative), each constrained by cannotRecommend arrays that structurally prohibit recommending "
    "what other personas are likely to recommend. The Claude Code Agent Teams feature (February 2026) was consequential here: "
    "before Agent Teams, VS Arena required simulated diversity through subagent calls in the same context window; after Agent "
    "Teams, each persona operates in its own independent context, producing substantively different outputs."
))

# 5.3
add_heading_apa(doc, "Human Checkpoint System: Five Required Cognitive Forcing Functions", level=2)
add_body(doc, (
    "Each REQUIRED checkpoint functions as a cognitive forcing function (Bucinca et al., 2021) that operationalizes the "
    "articulation method of cognitive apprenticeship: the researcher must verbalize a decision and its rationale before the "
    "system proceeds. The reduction from many checkpoints to five represents fading, a deliberate reduction of support."
))

add_body(doc, (
    "Automation bias is not primarily a knowledge problem; researchers often know they should evaluate AI recommendations "
    "critically. It is a cognitive habit problem: under time pressure and with confident-sounding AI output, critical evaluation "
    "is the path of most resistance. Cognitive forcing functions (Bucinca et al., 2021) are deliberate design elements that "
    "interrupt automatic processing and require effortful evaluation."
))

add_body(doc, (
    "The checkpoint design was itself iterative and marked by instructive failures. In v5.0 (January 25, 2026), the Sisyphus "
    "Protocol was introduced, making AI autonomous between checkpoints; it was removed the same day when it became clear that "
    "the protocol eliminated human control. The v6.0 Clean Slate Edition established the foundational principle: AI works "
    "between checkpoints; humans decide at them. Before the Claude Code hooks system (June 2025), checkpoints were behavioral "
    "instructions only; after hooks, they became architectural enforcement through PreToolUse interception."
))

add_body(doc, (
    "A particularly instructive failure occurred in v11.0: a silent bug in which all checkpoints auto-approved without user "
    "input. The system appeared to have oversight, but it did not. The bug was detected through architectural review, not user "
    "testing, illustrating how difficult genuine human oversight is to implement in autonomous AI systems. The fix involved "
    "SQLite-based verification of halt behavior."
))

add_body(doc, (
    "The five gate positions correspond to genuine cross-agent dependency points: decisions made at these points propagate "
    "through all subsequent agent outputs. Five is a judgment call, not an empirically derived optimum. The design case "
    "documents this honestly: the number reflects the designer's best current assessment of where oversight adds more value "
    "than it subtracts in cognitive load. Whether five checkpoints produce deliberation or habituated click-through is an "
    "empirical question for the companion evaluation paper."
))

# Figure 5: Checkpoint Flow
add_figure(doc,
    os.path.join(FIG_DIR, "figure5_checkpoint_flow.png"),
    "Human Checkpoint Flow: Five Required Cognitive Forcing Function Gates in the Research Design Workflow, With Cognitive Apprenticeship Method Labels",
    5
)

# ============================================================
# SECTION 6: DESIGN TENSIONS
# ============================================================
add_heading_apa(doc, "Design Tensions", level=1)

add_body(doc, (
    "Design cases in learning design often document what was built and why; the best ones also document what remains uncertain "
    "and why it matters. The following tensions are carried forward from the design decisions as open questions."
))

add_heading_apa(doc, "Simplicity Versus Specialization", level=2)
add_body(doc, (
    "Twenty-four agents in nine categories is more principled than 44, but it is still more complex than ChatGPT. For novice "
    "researchers, the 9-category architecture may itself be a barrier: knowing which category of agent to invoke requires prior "
    "knowledge of the research design process. The design bets that the right users (doctoral students in research design "
    "contexts) have enough process knowledge to navigate the architecture. This bet has not been tested."
))

add_heading_apa(doc, "Diversity Versus Coherence", level=2)
add_body(doc, (
    "The VS Arena generates five epistemologically divergent perspectives on a methodology question. For experienced researchers, "
    "this plurality is a resource. For novice researchers, it may be a source of confusion rather than illumination. No synthesis "
    "mechanism currently exists; the researcher must integrate competing perspectives without AI support. This is the most "
    "important unresolved design problem in Diverga's current state."
))

add_heading_apa(doc, "Oversight Rigor Versus Workflow Fluidity", level=2)
add_body(doc, (
    "Five checkpoints interrupt cognitive flow. The design trades fluidity for genuine oversight, but the exchange rate is "
    "unknown. Checkpoint fatigue is not hypothetical: in the pre-fix v11.0, the silent auto-approval behavior may have emerged "
    "partly because the original checkpoint count was too high and was quietly engineered around. The companion evaluation "
    "paper will investigate whether five REQUIRED checkpoints produce the intended deliberative behavior or habituated click-through."
))

# Table 4: Design Tensions
add_apa_table(doc,
    headers=["Tension", "Design Decision", "What Was Traded", "Open Question"],
    rows=[
        ["Simplicity vs. Specialization", "Multi-Agent Architecture", "User simplicity for domain coverage", "Is 24 agents still too many for novices?"],
        ["Diversity vs. Coherence", "VS and T-Score", "Convergence for epistemic plurality", "Do novices need synthesis support?"],
        ["Oversight vs. Fluidity", "Human Checkpoints", "Workflow fluidity for human control", "Do 5 checkpoints produce deliberation or click-through?"],
    ],
    table_num=4,
    title="Design Tensions Summary"
)

# ============================================================
# SECTION 7: REFLECTION
# ============================================================
add_heading_apa(doc, "Reflection and Future Directions", level=1)

add_heading_apa(doc, "What This Design Story Reveals", level=2)
add_body(doc, (
    "Learning design principles are not limited to explicit instructional settings. The decisions documented in this paper "
    "(structural diversity enforcement, mandatory cognitive interruption, typicality transparency) are recognizable to "
    "instructional designers even when they appear in an AI tool's architecture. The \"learning design in disguise\" framing "
    "is descriptive, not rhetorical: Diverga shapes researcher behavior through its form, not through explicit instruction, "
    "analogous to how a well-designed problem-based learning environment shapes student inquiry through task structure rather "
    "than direct teaching."
))

add_body(doc, (
    "The co-evolution with the Claude Code platform illustrates a broader phenomenon: AI tool designers are not designing on "
    "stable ground. Platform capabilities emerge and disappear; designs must adapt or become obsolete. The 41-day lifespan of "
    "cross-platform support and the architectural shift enabled by hooks and Agent Teams are not incidental details; they are "
    "defining features of what it means to design educational AI tools in a rapidly evolving ecosystem."
))

add_heading_apa(doc, "Transferable Design Principles", level=2)
add_body(doc, (
    "Three principles transfer to other AI tool designs for knowledge work. First, structural enforcement over instructional "
    "exhortation: design diversity into the system's architecture, not into the system's instructions. Second, transparency as "
    "metacognitive prompt: surface the information the user needs to calibrate their reliance on the system. Third, checkpoints "
    "as cognitive forcing functions: design mandatory pause points at consequential decision nodes, not at every step."
))

add_heading_apa(doc, "Connection to Companion Papers", level=2)
add_body(doc, (
    "This design case occupies Educational Design Research phases 1 and 2. The empirical evaluation of whether Diverga's design "
    "decisions produce the intended effects on researcher performance is the subject of a companion experimental paper (Paper 3). "
    "A separate methodological paper (Paper 2) addresses how to study artificial agency in multi-agent AI educational systems, "
    "using Diverga as the exemplar case. Authentic learning (Herrington, Reeves, & Oliver, 2010), which shares characteristics "
    "with Diverga's design context, will serve as the primary theoretical framework for Papers 2 and 3. The design case "
    "establishes the design rationale that the companion papers can reference without re-documenting the full system."
))

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_apa(doc, "References", level=1)

refs = [
    'Boling, E. (2010). The need for design cases: Disseminating design knowledge. International Journal of Designs for Learning, 1(1), 1-8.',
    'Bucinca, Z., Malaya, M. B., & Gajos, K. Z. (2021). To trust or to think: Cognitive forcing functions can reduce overreliance on AI in AI-assisted decision making. Proceedings of the ACM on Human-Computer Interaction, 5(CSCW1), Article 188.',
    'Collins, A., Brown, J. S., & Holum, A. (1991). Cognitive apprenticeship: Making thinking visible. American Educator, 15(3), 6-11.',
    'Collins, A., Brown, J. S., & Newman, S. E. (1989). Cognitive apprenticeship: Teaching the crafts of reading, writing, and mathematics. In L. B. Resnick (Ed.), Knowing, learning, and instruction: Essays in honor of Robert Glaser (pp. 453-494). Lawrence Erlbaum Associates.',
    'Herrington, J., Reeves, T. C., & Oliver, R. (2010). A guide to authentic e-learning. Routledge.',
    'Li, Z., Zhang, Y., & Zhao, H. (2025). More agents is less: Scaling properties of multi-agent LLM systems. arXiv:2602.03794.',
    'Markauskaite, L., & Goodyear, P. (2017). Epistemic fluency and professional education: Innovation, knowledgeable action and actionable knowledge. Springer.',
    'McKenney, S., & Reeves, T. C. (2012). Conducting educational design research. Routledge.',
    'Smith, K. M. (2010). Producing the rigorous design case. International Journal of Designs for Learning, 1(1), 9-20.',
    'Wang, T., Yu, P., & Zha, D. (2024). Mixture-of-agents yields emergent collaborative intelligence. Together AI Technical Report.',
    'Yang, M., & Harbor, J. (2023). Authentic learning design failures: The need for learner and contextual analysis and participatory design. International Journal of Designs for Learning, 14(1), 88-105.',
    'Zhao, Y., Huang, J., & Chen, X. (2024). Verbalized sampling for language model selection. arXiv:2510.01171.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"Document saved to: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT) / 1024:.0f} KB")

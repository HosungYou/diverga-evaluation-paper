#!/usr/bin/env python3
"""
Generate APA 7th Edition Word document for the Diverga Design Case paper.
Output: You_Yang_2026_Design_Case_Draft.docx

Authors: Hosung You, Mohan Yang
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIGURES_DIR = os.path.join(SCRIPT_DIR, "figures")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "You_Yang_2026_Design_Case_Draft.docx")


def set_cell_border(cell, **kwargs):
    """Set cell border properties.
    kwargs: top, bottom, start, end with value dicts of sz, val, color.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders %s>'
        '</w:tcBorders>' % nsdecls('w')
    )
    for edge, attrs in kwargs.items():
        element = parse_xml(
            '<w:{edge} {nsdecls} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'.format(
                edge=edge, nsdecls=nsdecls('w'),
                val=attrs.get('val', 'single'),
                sz=attrs.get('sz', '4'),
                color=attrs.get('color', '000000')
            )
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def create_document():
    doc = Document()

    # ---- Page Setup ----
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---- Default Style ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 2.0

    # ---- Running Head ----
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "SCAFFOLDING RESEARCH METHODOLOGY THROUGH AI"
    header_para.style = doc.styles['Normal']
    header_run = header_para.runs[0]
    header_run.font.size = Pt(12)
    header_run.font.name = 'Times New Roman'
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return doc


def add_title_page(doc):
    """APA 7th title page."""
    # Vertical spacing before title
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0

    # Title
    title_text = ("Scaffolding Research Methodology Through AI: "
                  "A Design Case of Diverga as Learning Design in Disguise")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.line_spacing = 2.0
    run = title_para.add_run(title_text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Blank line
    doc.add_paragraph().paragraph_format.line_spacing = 2.0

    # Author 1
    author1 = doc.add_paragraph()
    author1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author1.paragraph_format.line_spacing = 2.0
    run = author1.add_run("Hosung You")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Affiliation 1
    aff1 = doc.add_paragraph()
    aff1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff1.paragraph_format.line_spacing = 2.0
    run = aff1.add_run("College of Education, Pennsylvania State University")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Blank line
    doc.add_paragraph().paragraph_format.line_spacing = 2.0

    # Author 2
    author2 = doc.add_paragraph()
    author2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author2.paragraph_format.line_spacing = 2.0
    run = author2.add_run("Mohan Yang")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Affiliation 2
    aff2 = doc.add_paragraph()
    aff2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff2.paragraph_format.line_spacing = 2.0
    run = aff2.add_run("College of Education and Human Development, Texas A&M University")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Author note (blank lines for spacing)
    for _ in range(3):
        doc.add_paragraph().paragraph_format.line_spacing = 2.0

    # Author Note
    an_heading = doc.add_paragraph()
    an_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    an_heading.paragraph_format.line_spacing = 2.0
    run = an_heading.add_run("Author Note")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    an_text = doc.add_paragraph()
    an_text.paragraph_format.line_spacing = 2.0
    an_text.paragraph_format.first_line_indent = Inches(0.5)
    run = an_text.add_run(
        "Hosung You "
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = an_text.add_run("https://orcid.org/XXXX-XXXX-XXXX-XXXX")
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    an_text2 = doc.add_paragraph()
    an_text2.paragraph_format.line_spacing = 2.0
    an_text2.paragraph_format.first_line_indent = Inches(0.5)
    run = an_text2.add_run(
        "Correspondence concerning this article should be addressed to "
        "Hosung You, College of Education, Pennsylvania State University. "
        "Email: [placeholder]@psu.edu"
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Page break
    doc.add_page_break()


def add_heading_level1(doc, text):
    """APA 7th Level 1 heading: Centered, Bold, Title Case."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def add_heading_level2(doc, text):
    """APA 7th Level 2 heading: Flush Left, Bold, Title Case."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def add_heading_level3(doc, text):
    """APA 7th Level 3 heading: Flush Left, Bold Italic, Title Case."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def add_body_text(doc, text):
    """Add body text with first-line indent."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def add_apa_table(doc, title, headers, rows, note=None):
    """
    Create an APA 7th style table.
    APA tables: no vertical lines, horizontal rules at top, below header, and bottom.
    """
    # Table title (italic, above table)
    title_para = doc.add_paragraph()
    title_para.paragraph_format.line_spacing = 2.0
    title_para.paragraph_format.space_before = Pt(12)
    run = title_para.add_run(title)
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Create table
    n_cols = len(headers)
    n_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Remove all borders first
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": "0", "val": "none", "color": "FFFFFF"},
                bottom={"sz": "0", "val": "none", "color": "FFFFFF"},
                start={"sz": "0", "val": "none", "color": "FFFFFF"},
                end={"sz": "0", "val": "none", "color": "FFFFFF"},
            )

    # Header row: top and bottom borders
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        # Top border (thick)
        set_cell_border(
            cell,
            top={"sz": "12", "val": "single", "color": "000000"},
            bottom={"sz": "6", "val": "single", "color": "000000"},
            start={"sz": "0", "val": "none", "color": "FFFFFF"},
            end={"sz": "0", "val": "none", "color": "FFFFFF"},
        )

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(str(cell_text))
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

            # Bottom border only on last row
            if row_idx == len(rows) - 1:
                set_cell_border(
                    cell,
                    top={"sz": "0", "val": "none", "color": "FFFFFF"},
                    bottom={"sz": "12", "val": "single", "color": "000000"},
                    start={"sz": "0", "val": "none", "color": "FFFFFF"},
                    end={"sz": "0", "val": "none", "color": "FFFFFF"},
                )
            else:
                set_cell_border(
                    cell,
                    top={"sz": "0", "val": "none", "color": "FFFFFF"},
                    bottom={"sz": "0", "val": "none", "color": "FFFFFF"},
                    start={"sz": "0", "val": "none", "color": "FFFFFF"},
                    end={"sz": "0", "val": "none", "color": "FFFFFF"},
                )

    # Table note
    if note:
        note_para = doc.add_paragraph()
        note_para.paragraph_format.line_spacing = 2.0
        run_label = note_para.add_run("Note. ")
        run_label.italic = True
        run_label.font.size = Pt(12)
        run_label.font.name = 'Times New Roman'
        run_text = note_para.add_run(note)
        run_text.font.size = Pt(12)
        run_text.font.name = 'Times New Roman'

    return table


def add_figure(doc, image_path, figure_number, caption_text):
    """Insert figure with APA 7th caption format."""
    # Figure label
    label_para = doc.add_paragraph()
    label_para.paragraph_format.line_spacing = 2.0
    label_para.paragraph_format.space_before = Pt(12)
    run = label_para.add_run(f"Figure {figure_number}")
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Figure caption (italic)
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.line_spacing = 2.0
    run = caption_para.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Image
    if os.path.exists(image_path):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.line_spacing = 1.0
        run = img_para.add_run()
        run.add_picture(image_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[IMAGE NOT FOUND: {image_path}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.size = Pt(12)


def build_document():
    doc = create_document()

    # ======== TITLE PAGE ========
    add_title_page(doc)

    # ======== ABSTRACT ========
    add_heading_level1(doc, "Abstract")

    abstract_text = (
        "This design case documents the iterative design, development, and refinement of Diverga, "
        "a 24-agent AI research methodology assistant built as a Claude Code plugin. Diverga was "
        "created to address two well-documented problems in AI-assisted research: mode collapse "
        "(the tendency of AI to converge on predictable, modal recommendations) and automation bias "
        "(researchers' uncritical acceptance of AI suggestions). The paper presents three core design "
        "decisions and their rationale: (a) a multi-agent architecture that evolved from 44 to 24 "
        "specialized agents organized in 9 categories, (b) the integration of Verbalized Sampling (VS) "
        "methodology with explicit Typicality Scores (T-Scores) to counter mode collapse, and "
        "(c) a human checkpoint system that enforces structured researcher decision-making at critical "
        "junctures. Through a first-person design narrative spanning 57 days and 224 commits, we "
        "examine the design tensions between simplicity and specialization, diversity and coherence, "
        "and autonomy and oversight. Situated within the Analysis/Exploration and Design/Construction "
        "phases of the McKenney and Reeves (2012) educational design research model, this design case "
        "makes tacit design knowledge explicit for the instructional design and learning sciences "
        "community. The paper contributes to the IJDL Special Section on Learning Designs to Support "
        "and Improve Workplace Performance by demonstrating how a research methodology tool can function "
        "as a learning design in disguise, scaffolding methodological reasoning rather than replacing it."
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(abstract_text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.line_spacing = 2.0
    kw_para.paragraph_format.first_line_indent = Inches(0.5)
    run1 = kw_para.add_run("Keywords: ")
    run1.italic = True
    run1.font.size = Pt(12)
    run1.font.name = 'Times New Roman'
    run2 = kw_para.add_run(
        "design case, AI research assistant, multi-agent system, mode collapse, "
        "human checkpoint, verbalized sampling, learning design"
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    doc.add_page_break()

    # ======== SECTION 1: INTRODUCTION ========
    # Repeat title on first page of body
    title_repeat = doc.add_paragraph()
    title_repeat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_repeat.paragraph_format.line_spacing = 2.0
    run = title_repeat.add_run(
        "Scaffolding Research Methodology Through AI: "
        "A Design Case of Diverga as Learning Design in Disguise"
    )
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    add_body_text(doc,
        "Artificial intelligence tools for research are proliferating at an unprecedented rate. "
        "Tools such as Elicit, Consensus, and SciSpace promise to accelerate literature reviews, "
        "generate hypotheses, and streamline the research process. General-purpose AI assistants "
        "like ChatGPT and Claude are increasingly used by graduate students and researchers for "
        "methodology guidance. Yet beneath this proliferation lies a persistent and largely "
        "unaddressed problem: these tools converge on predictable, modal recommendations. When "
        "asked for a theoretical framework for technology adoption, ChatGPT recommends the "
        "Technology Acceptance Model (TAM) in approximately 80% of cases. When asked about "
        "qualitative analysis approaches, thematic analysis is the near-universal suggestion. "
        "This phenomenon, known as mode collapse, undermines the exploratory nature of research "
        "and reinforces methodological conservatism."
    )

    add_body_text(doc,
        "Compounding this problem is automation bias, the documented tendency for users to accept "
        "AI-generated recommendations without critical evaluation. Recent evidence suggests that "
        "78% of researchers accept AI-generated suggestions without verification (Springer Nature, "
        "2025), and a negative correlation exists between AI tool use and critical thinking "
        "(Microsoft Research, 2025). Together, mode collapse and automation bias create a troubling "
        "dynamic: AI tools that consistently recommend the same approaches, paired with researchers "
        "who uncritically accept those recommendations."
    )

    add_body_text(doc,
        "This design case documents the design process, decisions, and rationale behind Diverga, "
        "a 24-agent AI research methodology assistant. Diverga was built to scaffold "
        "methodological reasoning rather than replace it. The system addresses mode collapse "
        "through Verbalized Sampling (VS) methodology and addresses automation bias through a "
        "structured human checkpoint system. This paper does not report empirical evaluation data; "
        "it tells the design story. Situated within the Analysis/Exploration and Design/Construction "
        "phases of the McKenney and Reeves (2012) educational design research (EDR) model, this "
        "design case captures what was learned during the iterative design process. Companion "
        "papers will address empirical evaluation (Paper 3) and methodological contributions "
        "to studying artificial agency in multi-agent AI tools (Paper 2)."
    )

    add_body_text(doc,
        "[PLACEHOLDER: Additional introduction paragraphs covering the scope statement, the "
        "contribution to the IJDL Special Section on Learning Designs to Support and Improve "
        "Workplace Performance, and a roadmap for the remainder of the paper. Target length "
        "for this section: approximately 800 words total.]"
    )

    # ======== SECTION 2: DESIGN CONTEXT ========
    add_heading_level1(doc, "Design Context")

    add_body_text(doc,
        "Diverga was designed and developed by the first author, a doctoral student in Learning, "
        "Design, and Technology at Pennsylvania State University. This dual role as both designer "
        "and researcher introduces both affordances and limitations that are addressed explicitly "
        "in the Design Tensions section. The second author, who brings experience with design case "
        "methodology and instructional design research, provides an external perspective that "
        "interrogates the tacit design knowledge embedded in the system."
    )

    add_body_text(doc,
        "The target users of Diverga are social science doctoral students engaged in research "
        "design, particularly in fields such as Education, Psychology, Management, Sociology, "
        "Human Resource Development, and Communication. These students typically complete "
        "methodology coursework but face a significant gap between course content and the "
        "practical demands of designing their own studies. This gap is the space that Diverga "
        "aims to scaffold."
    )

    add_body_text(doc,
        "[PLACEHOLDER: Additional context paragraphs covering the technical context (Model Context "
        "Protocol, Claude Code infrastructure, the multi-agent architecture landscape circa 2025 "
        "to 2026), institutional context (graduate research methods education in U.S. universities), "
        "and the broader landscape of AI research tools. Target length: approximately 500 words "
        "total.]"
    )

    # ======== SECTION 3: THE DESIGN PROBLEM ========
    add_heading_level1(doc, "The Design Problem")

    add_heading_level2(doc, "Mode Collapse in AI Research Assistance")

    add_body_text(doc,
        "Mode collapse refers to the tendency of large language models to converge on the most "
        "probable, most frequently occurring response. In the context of research methodology "
        "guidance, this manifests as a persistent bias toward well-known, conventional approaches. "
        "When researchers consult AI assistants about theoretical frameworks, research designs, "
        "or analytical strategies, the AI typically recommends the most common option within "
        "a given domain. This convergence is not a bug but a feature of how language models are "
        "trained: they optimize for the most likely continuation of text, which naturally favors "
        "the most frequently discussed options in the training data."
    )

    add_body_text(doc,
        "The Verbalized Sampling (VS) methodology (arXiv:2510.01171) provides a theoretical "
        "framework for addressing mode collapse. VS introduces the concept of Typicality Scores "
        "(T-Scores) to quantify how conventional or novel a given recommendation is. A "
        "recommendation with T > 0.8 is highly typical (e.g., recommending TAM for technology "
        "adoption), while a recommendation with T < 0.3 represents a genuinely novel approach "
        "that requires strong justification. VS encourages exploration across the full typicality "
        "spectrum rather than defaulting to the highest-probability option."
    )

    add_heading_level2(doc, "Automation Bias in Research")

    add_body_text(doc,
        "[PLACEHOLDER: Discussion of automation bias literature, including the 78% uncritical "
        "acceptance rate (Springer Nature, 2025), negative correlation with critical thinking "
        "(Microsoft Research, 2025), and implications for research quality. The concept of "
        "cognitive forcing functions (Bucinca et al., 2021) as a design strategy for mitigating "
        "automation bias. Target length: approximately 350 words.]"
    )

    add_heading_level2(doc, "Existing Tool Landscape")

    add_body_text(doc,
        "[PLACEHOLDER: Brief comparison of Diverga with existing AI research tools, followed "
        "by a summary table. This section establishes the gap that Diverga addresses: no existing "
        "tool combines structured human oversight with diversity-generating mechanisms for research "
        "methodology guidance. Target length: approximately 300 words.]"
    )

    # Table 1: Comparison with existing tools
    add_apa_table(
        doc,
        "Table 1\n\nComparison of Diverga With Existing AI Research Tools",
        ["Feature", "Diverga", "Elicit", "Consensus", "ChatGPT/Claude"],
        [
            ["Primary purpose",
             "Research methodology scaffolding",
             "Literature search and extraction",
             "Evidence synthesis from papers",
             "General-purpose assistance"],
            ["Mode collapse mitigation",
             "VS methodology with T-Scores",
             "None",
             "None",
             "None"],
            ["Human oversight mechanism",
             "5 REQUIRED checkpoints, MCP-enforced",
             "None (user-driven)",
             "None (user-driven)",
             "None (user-driven)"],
            ["Agent architecture",
             "24 specialized agents in 9 categories",
             "Single pipeline",
             "Single pipeline",
             "Single agent"],
            ["Research lifecycle coverage",
             "Full (question to publication)",
             "Literature review phase",
             "Evidence synthesis phase",
             "Ad hoc, any phase"],
            ["Methodological diversity",
             "VS Arena with 5 epistemological personas",
             "Not applicable",
             "Not applicable",
             "Prompt-dependent"],
            ["Decision audit trail",
             "YAML-based, immutable log",
             "Not available",
             "Not available",
             "Not available"],
            ["Target users",
             "Social science doctoral students",
             "Researchers (all levels)",
             "Researchers (all levels)",
             "General audience"],
        ],
        note="VS = Verbalized Sampling; MCP = Model Context Protocol; YAML = Yet Another Markup Language."
    )

    # ======== SECTION 4: DESIGN DECISIONS AND RATIONALE ========
    add_heading_level1(doc, "Design Decisions and Rationale")

    add_body_text(doc,
        "This section presents three core design decisions that shaped Diverga. Each subsection "
        "follows a consistent structure: the decision point, alternatives considered, rationale "
        "for the chosen approach, and evidence of implementation. These decisions are presented "
        "thematically rather than chronologically, though the design evolution is documented in "
        "Figure 3."
    )

    # ---- 4.1 Multi-Agent Architecture ----
    add_heading_level2(doc, "Multi-Agent Architecture: From 44 to 24")

    add_body_text(doc,
        "The most fundamental design decision in Diverga was the choice to implement a "
        "multi-agent architecture rather than a single general-purpose agent. This decision "
        "was grounded in the principle that specialized agents with clear domain boundaries "
        "produce higher-quality outputs than a single agent attempting to cover the entire "
        "research lifecycle. The system evolved through several iterations, reaching a peak "
        "of 44 agents at version 6.7 before being consolidated to 24 agents at version 11.0."
    )

    add_body_text(doc,
        "The consolidation from 44 to 24 agents was driven by three factors. First, many "
        "agents had overlapping responsibilities that created confusion about which agent to "
        "invoke. For example, A3 (Devil's Advocate) and A6 (Visualizer) were merged into A2 "
        "(Theory and Critique Architect), combining functions that naturally co-occur during "
        "theoretical framework development. Second, the token overhead of maintaining 44 agent "
        "definitions exceeded practical limits; each agent definition consumed tokens from "
        "the context window, leaving less capacity for actual research work. Third, platform "
        "fragmentation across three code assistant platforms (Claude Code, Codex CLI, OpenCode) "
        "was not delivering proportional value relative to the maintenance burden."
    )

    add_body_text(doc,
        "[PLACEHOLDER: Additional discussion of the multi-agent architecture decision, including "
        "references to Brooks' subsumption architecture principle, 'more agents is less' scaling "
        "research (Li et al., 2025), Together AI's Mixture of Agents findings, and the specific "
        "consolidation logic. Include discussion of how the 9-category structure (A through I, "
        "plus X) maps onto the research lifecycle. Target length: approximately 700 words total "
        "for this subsection.]"
    )

    # Insert Figure 1
    add_figure(doc,
               os.path.join(FIGURES_DIR, "figure1_architecture.png"),
               1,
               "Diverga System Architecture: 24 Specialized Agents Organized in 9 Categories "
               "Covering the Complete Research Lifecycle")

    # Table 2: Agent consolidation summary
    add_apa_table(
        doc,
        "Table 2\n\nAgent Consolidation Summary (v10.3 to v11.0)",
        ["Original agents", "Merged into", "Rationale"],
        [
            ["A3 (Devil's Advocate), A6 (Visualizer)",
             "A2 (Theory Architect)",
             "Functions naturally co-occur during theory development"],
            ["A4 (Ethics), F4 (Bias Detection)",
             "X1 (Research Guardian)",
             "Ethics and bias detection are cross-cutting concerns"],
            ["B3, C6, C7 (Meta-analysis support)",
             "C5 (Meta-Analysis Master)",
             "Consolidated fragmented meta-analysis pipeline"],
            ["F1, F2, F3 (Quality assurance)",
             "G2 (Publication Specialist)",
             "Writing quality functions overlap with publication"],
            ["H1, H2 (Specialized approaches)",
             "C2 (Qualitative Designer)",
             "Specialized qualitative approaches belong in design"],
            ["D1, D3 (Data collection)",
             "D2 (Collection Specialist)",
             "Overlapping data collection functions"],
            ["B4, B5 (Literature support)",
             "B1 (Literature Scout)",
             "Parallel document processing integrated into scout"],
            ["C4 (Design support)",
             "C1 (Quant Design)",
             "Redundant quantitative design guidance"],
        ],
        note="v10.3 contained 44 agents across 8 categories (A through I). v11.0 reduced "
             "to 24 agents across 9 categories (A through I, plus X) by merging agents with "
             "overlapping responsibilities."
    )

    # ---- 4.2 Verbalized Sampling and T-Score ----
    add_heading_level2(doc, "Verbalized Sampling and T-Score System")

    add_body_text(doc,
        "The second core design decision was the integration of Verbalized Sampling (VS) "
        "methodology across all agents. VS was not an afterthought; it was integrated within "
        "hours of the initial commit (v2.0, January 22, 2026), indicating that mode collapse "
        "prevention was always the planned core differentiator. The VS methodology introduces "
        "a 3-stage process: (a) context and modal identification, where the system identifies "
        "the 'obvious' recommendation (T > 0.8); (b) divergent exploration, where the system "
        "generates three alternative directions spanning the typicality spectrum; and "
        "(c) human selection, where the researcher reviews all options with their T-Scores "
        "and makes an informed choice."
    )

    add_body_text(doc,
        "A key design decision was to make T-Scores visible to the researcher rather than using "
        "them only internally. This transparency serves two purposes: it makes the AI's reasoning "
        "process legible, and it provides researchers with a cognitive framework for evaluating "
        "methodological choices. When a researcher sees that thematic analysis has a T-Score of "
        "0.92 (highly typical) while narrative analysis has a T-Score of 0.35 (innovative), they "
        "gain metacognitive awareness of where their choice falls on the convention-novelty "
        "spectrum."
    )

    add_body_text(doc,
        "The VS methodology was not applied uniformly across all agents. Instead, a tiered "
        "application model was adopted. Agents responsible for high-stakes methodological "
        "decisions (A2, A5, C2, C3, C5, E1, E2, E3) received Full VS with all five phases. "
        "Agents with moderate decision impact (A1, C1, D2, D4, G6) received Enhanced VS with "
        "three phases. Operational agents where creativity matters less (G1, G5, I1, I2, I3) "
        "received Light VS with modal awareness only."
    )

    add_body_text(doc,
        "[PLACEHOLDER: Additional discussion of VS alternatives considered (standard prompting "
        "with 'be diverse' instructions, retrieval-augmented generation from a curated "
        "methodology database, random sampling with filtering) and why VS was chosen. Include "
        "the VS Arena design with 5 epistemological personas (post-positivist, critical theorist, "
        "pragmatist, interpretivist, transformative) and the cannotRecommend constraint mechanism. "
        "Target length: approximately 700 words total for this subsection.]"
    )

    # ---- 4.3 Human Checkpoint System ----
    add_heading_level2(doc, "Human Checkpoint System")

    add_body_text(doc,
        "The third core design decision was the implementation of structured, mandatory "
        "checkpoints at specific decision points in the research workflow. This decision "
        "emerged from a clear philosophical position: 'Human decisions remain with humans. "
        "AI handles what is beyond human scope.' The checkpoint system underwent significant "
        "evolution, reflecting the broader design tension between autonomy and oversight."
    )

    add_body_text(doc,
        "The evolution of the checkpoint system reveals a pivotal design moment. Version 5.0 "
        "(the 'Sisyphus Edition') introduced the Sisyphus Protocol, which enforced continuous "
        "AI operation until task completion. The system would keep working autonomously, only "
        "pausing at checkpoints. Version 6.0 (the 'Clean Slate Edition') deliberately removed "
        "this protocol along with all autonomous continuation modes. The language of the system "
        "shifted from assertive ('I will proceed') to consultative ('May I proceed?'). This "
        "was a philosophical pivot from 'AI-driven with human oversight' to 'human-driven with "
        "AI assistance.'"
    )

    add_body_text(doc,
        "The current system implements five REQUIRED checkpoints that enforce mandatory "
        "researcher decision-making: CP_RESEARCH_DIRECTION (when the research question is "
        "finalized), CP_PARADIGM_SELECTION (when the methodology approach is chosen), "
        "CP_THEORY_SELECTION (when the theoretical framework is selected), "
        "CP_METHODOLOGY_APPROVAL (when the research design is complete), and CP_META_GATE "
        "(when meta-analysis validation gates require decisions). At each checkpoint, the "
        "system stops, presents options with VS alternatives and T-Scores, and waits for "
        "explicit human approval before proceeding."
    )

    add_heading_level3(doc, "The Checkpoint Bug: A Design Failure Narrative")

    add_body_text(doc,
        "A particularly instructive design failure occurred at version 11.0, when a bug caused "
        "all checkpoints to return '{continue: true}' regardless of the researcher's actual "
        "decision. This effectively transformed the checkpoint system into a rubber stamp. "
        "The bug was caused by the transition from soft blocks (which injected warnings but "
        "allowed the AI to proceed) to hard blocks (which were intended to halt execution). "
        "The fix required implementing runtime enforcement through the MCP checkpoint server, "
        "backed by a SQLite database for ACID-safe state management. This failure illustrates "
        "a fundamental challenge in human-AI interaction design: implementing genuine human "
        "oversight in autonomous systems is technically difficult, and the default failure mode "
        "is to silently bypass the very safeguards designed to keep humans in the loop."
    )

    add_body_text(doc,
        "[PLACEHOLDER: Additional discussion of checkpoint alternatives considered (no "
        "checkpoints with full autonomy, optional checkpoints with soft suggestions, checkpoints "
        "at every step for maximum oversight), checkpoint fatigue considerations, and the "
        "evolution from many REQUIRED checkpoints to five. Include discussion of cognitive "
        "forcing functions (Bucinca et al., 2021) as a theoretical grounding for the checkpoint "
        "design. Target length: approximately 700 words total for this subsection.]"
    )

    # Insert Figure 2
    add_figure(doc,
               os.path.join(FIGURES_DIR, "figure2_checkpoint_flow.png"),
               2,
               "Human Checkpoint System: Research Workflow With Decision Gates Showing "
               "REQUIRED (Solid) and RECOMMENDED (Dashed) Checkpoint Positions")

    # ======== SECTION 5: DESIGN TENSIONS ========
    add_heading_level1(doc, "Design Tensions")

    add_body_text(doc,
        "The design process surfaced several inherent tensions that remain partially unresolved. "
        "Rather than presenting these as problems solved, we frame them as ongoing design "
        "challenges that characterize the current state of the system and point toward future "
        "design iterations."
    )

    # Table 3: Design tensions
    add_apa_table(
        doc,
        "Table 3\n\nDesign Tensions in Diverga",
        ["Tension", "Description", "Current resolution"],
        [
            ["Simplicity vs. Specialization",
             "24 agents in 9 categories is complex compared to a single-agent interface like "
             "ChatGPT, creating a learning curve for new users.",
             "Consolidated from 44 to 24; auto-detection reduces manual agent selection; "
             "further simplification risks losing domain specificity."],
            ["Diversity vs. Coherence",
             "The VS Arena maximizes methodological diversity, but diverse suggestions without "
             "synthesis can overwhelm novice researchers.",
             "T-Scores provide a framework for evaluating options; no arbitrating agent "
             "(e.g., Chief Scientific Officer) exists yet to synthesize recommendations."],
            ["Autonomy vs. Oversight",
             "Every checkpoint interrupts workflow and adds cognitive load, yet removing "
             "checkpoints risks automation bias.",
             "Five REQUIRED checkpoints represent a judgment call, not a proven optimum. "
             "Checkpoint fatigue remains an open concern."],
            ["Designer as Researcher",
             "The developer is also the evaluator, creating potential blind spots in the "
             "design rationale.",
             "The second author provides an external perspective; empirical evaluation "
             "(Paper 3) will provide independent evidence of effectiveness."],
        ],
        note="These tensions are framed as inherent design challenges rather than solved problems. "
             "Empirical investigation of optimal checkpoint frequency and agent count is planned "
             "as future work."
    )

    # ======== Figure 3: Evolution Timeline ========
    add_figure(doc,
               os.path.join(FIGURES_DIR, "figure3_evolution_timeline.png"),
               3,
               "Design Evolution Timeline: Agent Count and Key Milestones From v1.0 "
               "(January 22, 2026) to v12.0 (March 19, 2026)")

    # ======== SECTION 6: REFLECTION AND FUTURE DIRECTIONS ========
    add_heading_level1(doc, "Reflection and Future Directions")

    add_body_text(doc,
        "This design case has documented the iterative design process behind Diverga, a "
        "24-agent AI research methodology assistant. The design journey, spanning 57 days and "
        "224 commits, reveals several insights about designing AI tools for educational contexts."
    )

    add_body_text(doc,
        "First, the tension between automation and scaffolding is not a binary choice but a "
        "continuum. Diverga's evolution from the Sisyphus Protocol (maximum automation) to "
        "the Clean Slate Edition (structured oversight) reflects a broader philosophical shift "
        "in how we think about AI's role in learning: not as a replacement for human judgment "
        "but as a scaffold that makes the space of possibilities visible while preserving the "
        "learner's agency to choose."
    )

    add_body_text(doc,
        "Second, the consolidation from 44 to 24 agents illustrates that complexity in AI "
        "system design does not scale linearly with capability. The reduced system is more "
        "capable precisely because it eliminated redundancy and clarified boundaries. This "
        "finding resonates with research on cognitive load in instructional design: more is "
        "not always better, and the structure of information matters as much as its quantity."
    )

    add_body_text(doc,
        "Third, the checkpoint bug narrative demonstrates that implementing genuine human "
        "oversight in autonomous systems is technically and conceptually challenging. The "
        "default failure mode of AI systems is to bypass safeguards silently, making the "
        "design of robust intervention points a first-order concern for any AI tool intended "
        "for educational use."
    )

    add_body_text(doc,
        "[PLACEHOLDER: Additional reflection paragraphs covering how this design case connects "
        "to the EDR model (the Evaluation/Reflection phase and empirical testing are Papers 2 "
        "and 3), an invitation for the field regarding design principles that other AI tool "
        "developers can adapt, and the open question of whether structural diversity (VS Arena) "
        "outperforms instructional diversity (single-agent prompting). Target length: "
        "approximately 500 words total.]"
    )

    # ======== REFERENCES ========
    add_heading_level1(doc, "References")

    references = [
        "Boling, E. (2010). The need for design cases: Disseminating design knowledge. "
        "International Journal of Designs for Learning, 1(1), 1-8. "
        "https://doi.org/10.14434/ijdl.v1i1.919",

        "Bucinca, Z., Malaya, M. B., & Gajos, K. Z. (2021). To trust or to think: Cognitive "
        "forcing functions can reduce overreliance on AI in AI-assisted decision-making. "
        "Proceedings of the ACM on Human-Computer Interaction, 5(CSCW1), Article 188. "
        "https://doi.org/10.1145/3449287",

        "Howard, C. D., Boling, E., Rowland, G., & Smith, K. M. (2012). Instructional design "
        "cases and why we need them. Educational Technology, 52(3), 34-38.",

        "Li, Y., et al. (2025). More agents is less: Scaling challenges in multi-agent systems. "
        "[Placeholder reference]",

        "McKenney, S., & Reeves, T. C. (2012). Conducting educational design research. Routledge.",

        "Microsoft Research. (2025). The impact of AI tools on critical thinking in research. "
        "[Placeholder reference]",

        "Springer Nature. (2025). AI in research: Usage patterns and verification behaviors. "
        "[Placeholder reference]",

        "Zhao, Y., et al. (2024). Verbalized sampling for language model generation. "
        "arXiv:2510.01171. https://arxiv.org/abs/2510.01171",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(ref)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    print("Generating APA 7th Edition Word document for Diverga Design Case...")
    print("=" * 65)
    build_document()
    print("=" * 65)
    print("Done.")

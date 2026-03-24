"""
generate_word_docs.py
Converts Paper 2 (BJET) and Paper 3 (ETR&D) markdown drafts to Word documents.
Applies em-dash removal and humanization before conversion.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────
# 1. HUMANIZATION ENGINE
# ─────────────────────────────────────────

def replace_emdash(text: str) -> str:
    """
    Contextual em-dash (—) replacement.
    Handles the five most common academic em-dash patterns.
    """
    # Pattern 1: X — characterized by Y —  →  X, characterized by Y,
    # e.g. "artificial agency — characterized by goal persistence — represents"
    text = re.sub(
        r'\s—\s(characterized by|defined by|marked by|distinguished by)',
        r', \1', text
    )
    text = re.sub(
        r'(characterized by[^—]+)\s—\s',
        r'\1, ', text
    )

    # Pattern 2: Paired em-dashes as parenthetical  — X —  →  (X)
    # Match: word/punct — content (no nested dash) — word/punct
    text = re.sub(
        r'\s—\s([^—]{1,80}?)\s—\s',
        r' (\1) ', text
    )

    # Pattern 3: "not X" or "not Y" after em-dash  → comma
    # e.g. "a form of intelligence — not a form of agency" → ", not a form"
    text = re.sub(r'\s—\s(not\b)', r', \1', text)

    # Pattern 4: Em-dash introducing a list or elaboration after colon-like context
    # "three criteria — interactivity, bounded autonomy, and adaptability"
    # → "three criteria: interactivity, bounded autonomy, and adaptability"
    text = re.sub(
        r'(criteria|mechanisms|features|elements|dimensions|aspects|components|properties|factors)\s—\s',
        r'\1: ', text
    )

    # Pattern 5: Em-dash between noun phrase and appositive description
    # "X — a methodology that..." → "X, a methodology that..."
    text = re.sub(r'\s—\s(a |an |the )', r', \1', text)

    # Pattern 6: Em-dash introducing a consequence or result clause
    # "X — revealing not merely..." → "X, revealing not merely..."
    text = re.sub(
        r'\s—\s(revealing|showing|demonstrating|indicating|suggesting|enabling|allowing|ensuring|requiring|providing|creating|generating|producing|making)',
        r', \1', text
    )

    # Pattern 7: Em-dash at end of sentence introducing elaboration
    # "...measurement challenge — systems characterized by..." → "...measurement challenge: systems..."
    text = re.sub(r'([a-z])\s—\s([A-Z])', r'\1: \2', text)

    # Pattern 8: Remaining solo em-dashes → semicolons or commas
    # If still has em-dash surrounded by lowercase: use semicolon
    text = re.sub(r'([a-z,])\s—\s([a-z])', r'\1; \2', text)
    # Any remaining em-dashes → comma
    text = text.replace(' — ', ', ')
    text = text.replace('—', ', ')

    return text


def humanize_text(text: str) -> str:
    """Remove common AI writing patterns beyond em-dashes."""

    # Fix: "It is important to note that" → trim
    text = re.sub(r'It is important to note that\s+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'It should be noted that\s+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'It is worth noting that\s+', '', text, flags=re.IGNORECASE)

    # Fix: "Notably," at sentence start → rephrase as nothing (let sentence stand)
    text = re.sub(r'^Notably,\s+', '', text, flags=re.MULTILINE)

    # Fix: "Importantly," at sentence start → keep if truly important, else remove
    # (light touch — just remove the word, keep sentence)
    text = re.sub(r'^Importantly,\s+', '', text, flags=re.MULTILINE)

    # Fix: "In this paper, we" → "This paper"
    text = re.sub(r'In this paper,?\s+we\s+', 'This paper ', text, flags=re.IGNORECASE)

    # Fix: "as previously mentioned" → remove
    text = re.sub(r',?\s*as (previously |already )?mentioned', '', text, flags=re.IGNORECASE)

    # Fix: overcapitalized "REQUIRED" → "required" in prose (not in tables)
    # Only in flowing prose (surrounded by spaces and lowercase)
    text = re.sub(r'\bREQUIRED\b', 'required', text)
    text = re.sub(r'\bADVISORY\b', 'advisory', text)

    # Fix: Excessive parenthetical stacking "(a)...(b)...(c)" mid-sentence
    # These are fine — leave them as is for academic prose

    # Clean up double spaces created by replacements
    text = re.sub(r'  +', ' ', text)
    text = re.sub(r' ,', ',', text)
    text = re.sub(r' ;', ';', text)
    text = re.sub(r'\( ', '(', text)
    text = re.sub(r' \)', ')', text)

    return text


def process_markdown(raw: str) -> str:
    """Apply em-dash removal and humanization to full document text."""
    result = replace_emdash(raw)
    result = humanize_text(result)
    return result


# ─────────────────────────────────────────
# 2. WORD DOCUMENT BUILDER
# ─────────────────────────────────────────

def set_font(run, name='Times New Roman', size_pt=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_styled_paragraph(doc, text, style='Normal', alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          bold=False, italic=False, size=12, space_before=0, space_after=6,
                          color=None):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size_pt=size, bold=bold, italic=italic, color=color)
    return p


def add_section_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size_pt=sizes.get(level, 11), bold=True)
    return p


def add_draft_notice(doc, notice_text):
    """Add a shaded draft notice box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3CD')
    pPr.append(shd)
    run = p.add_run(notice_text)
    set_font(run, size_pt=10, italic=True, color=(120, 80, 0))
    return p


def add_table_from_markdown(doc, header_row, data_rows):
    """Build a formatted Word table from parsed markdown table data."""
    n_cols = len(header_row)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header_row):
        hdr_cells[i].text = cell_text.strip()
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
        # Shade header
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8E8E8')
        tcPr.append(shd)

    # Data rows
    for row_data in data_rows:
        if not any(c.strip() for c in row_data):
            continue
        row = table.add_row()
        for i, cell_text in enumerate(row_data[:n_cols]):
            row.cells[i].text = cell_text.strip()
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacing after table


def parse_markdown_table(lines):
    """Parse markdown table lines into header + data rows."""
    header = None
    rows = []
    for line in lines:
        if line.startswith('|'):
            cells = [c for c in line.split('|')[1:-1]]
            # Skip separator rows
            if all(re.match(r'^[-: ]+$', c.strip()) for c in cells if c.strip()):
                continue
            if header is None:
                header = cells
            else:
                rows.append(cells)
    return header, rows


# ─────────────────────────────────────────
# 3. MARKDOWN → WORD PARSER
# ─────────────────────────────────────────

def markdown_to_word(md_text: str, output_path: str, paper_num: int):
    """Full markdown → Word conversion with em-dash humanization."""

    # Apply humanization
    md_text = process_markdown(md_text)

    doc = Document()

    # Page setup: 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(22)  # ~1.5 line spacing

    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_lines = []
    in_review_note = False

    while i < len(lines):
        line = lines[i]

        # Skip metadata lines (**, status lines at top)
        if line.startswith('**Target:**') or line.startswith('**Track:**') or \
           line.startswith('**Status:**') or line.startswith('**Word count') or \
           line.startswith('**Abstract deadline') or line.startswith('**Relationship:') or \
           line.startswith('**IRB status:') or line.startswith('**Data collection'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            set_font(run, size_pt=10, italic=True, color=(80, 80, 80))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Draft / collaborator note in blockquote
        if line.startswith('> **Note for Dr. Yang:**') or line.startswith('> **Note:**'):
            notice = line.lstrip('> ').replace('**Note for Dr. Yang:**', 'Note for Dr. Yang:').replace('**Note:**', 'Note:')
            # Collect continuation lines
            j = i + 1
            while j < len(lines) and lines[j].startswith('> '):
                notice += ' ' + lines[j].lstrip('> ')
                j += 1
            add_draft_notice(doc, notice)
            i = j
            continue

        # Review note (skip in final doc — these are internal)
        if '> **[REVIEW NOTE' in line:
            j = i + 1
            while j < len(lines) and lines[j].startswith('>'):
                j += 1
            i = j
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 70)
            set_font(run, size_pt=8, color=(180, 180, 180))
            i += 1
            continue

        # H1 title
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if 'DRAFT' in text or 'Paper 2:' in text or 'Paper 3:' in text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                set_font(run, size_pt=14, bold=True)
            i += 1
            continue

        # H2 section
        if line.startswith('## '):
            text = line[3:].strip()
            add_section_heading(doc, text, level=1)
            i += 1
            continue

        # H3 subsection
        if line.startswith('### '):
            text = line[4:].strip()
            add_section_heading(doc, text, level=2)
            i += 1
            continue

        # H4 sub-subsection
        if line.startswith('#### '):
            text = line[5:].strip()
            add_section_heading(doc, text, level=3)
            i += 1
            continue

        # Table detection
        if line.startswith('|') and not in_table:
            in_table = True
            table_lines = [line]
            i += 1
            continue

        if in_table:
            if line.startswith('|'):
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                header, rows = parse_markdown_table(table_lines)
                if header:
                    add_table_from_markdown(doc, header, rows)
                in_table = False
                table_lines = []
                continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # remove bold markers
            text = re.sub(r'\*(.+?)\*', r'\1', text)       # remove italic markers
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            set_font(run, size_pt=11)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            set_font(run, size_pt=11)
            i += 1
            continue

        # Figure/table caption lines
        if line.startswith('**Figure') or line.startswith('**Table') or \
           line.startswith('Caption:') or line.startswith('*Figure') or line.startswith('*Table'):
            text = re.sub(r'\*+', '', line).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            set_font(run, size_pt=10, italic=True)
            i += 1
            continue

        # [OUTLINE] tags and internal notes — include as gray text
        if '[OUTLINE' in line or '[REVIEW NOTE' in line or line.strip().startswith('[') and ']' in line:
            text = re.sub(r'\*+', '', line).strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                set_font(run, size_pt=10, italic=True, color=(130, 130, 130))
            i += 1
            continue

        # Inline bold/italic handling for regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0)

            # Parse inline bold/italic patterns
            segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', line)
            for seg in segments:
                if seg.startswith('**') and seg.endswith('**'):
                    run = p.add_run(seg[2:-2])
                    set_font(run, size_pt=11, bold=True)
                elif seg.startswith('*') and seg.endswith('*'):
                    run = p.add_run(seg[1:-1])
                    set_font(run, size_pt=11, italic=True)
                elif seg.startswith('`') and seg.endswith('`'):
                    run = p.add_run(seg[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    if seg:
                        run = p.add_run(seg)
                        set_font(run, size_pt=11)

        i += 1

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ─────────────────────────────────────────
# 4. MAIN
# ─────────────────────────────────────────

BASE = Path('/Users/newhosung/diverga-evaluation-paper')

papers = [
    {
        'input': BASE / 'Paper2_BJET_Methodological/paper2_proposal_draft_v1.md',
        'output': BASE / 'Paper2_BJET_Methodological/Paper2_BJET_proposal_draft_v1.docx',
        'num': 2,
    },
    {
        'input': BASE / 'Paper3_Experimental/paper3_etrd_proposal_draft_v1.md',
        'output': BASE / 'Paper3_Experimental/Paper3_ETRnD_proposal_draft_v1.docx',
        'num': 3,
    },
]

for paper in papers:
    raw = paper['input'].read_text(encoding='utf-8')
    print(f"\nProcessing Paper {paper['num']}...")

    # Show em-dash count before/after
    before = raw.count('—')
    processed = process_markdown(raw)
    after = processed.count('—')
    print(f"  Em-dashes: {before} → {after}")

    markdown_to_word(processed, str(paper['output']), paper['num'])

print("\nDone.")

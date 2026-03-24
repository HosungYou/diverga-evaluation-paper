"""
generate_word_apa7.py
APA 7th edition Word manuscript generator for Paper 2 (BJET) and Paper 3 (ETR&D).

Format:
  - Title page (anonymous): title, running head, word count, no author names
  - Abstract page (p.2): structured (BJET) / unstructured (ETR&D)
  - Body: double-spaced, Times New Roman 12pt, 0.5" first-line indent
  - Headings: APA Level 1 (bold centered), L2 (bold left), L3 (bold italic left)
  - Tables: APA format (top/header/bottom rules only, no vertical lines)
  - Figures: actual PNG images embedded with APA-style captions
  - [OUTLINE] sections: retained as draft markers in grey italic
  - Self-citations anonymized: [Author(s), year]
"""

import re
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path('/Users/newhosung/diverga-evaluation-paper')

# ─────────────────────────────────────────────────────
# 1.  TEXT PROCESSING
# ─────────────────────────────────────────────────────

def replace_emdash(text):
    text = re.sub(r'\s—\s(characterized by|defined by|marked by|distinguished by|not\b)', r', \1', text)
    text = re.sub(r'\s—\s([^—]{1,80}?)\s—\s', r' (\1) ', text)
    text = re.sub(r'(criteria|mechanisms|features|elements|dimensions|aspects|components|properties|factors)\s—\s', r'\1: ', text)
    text = re.sub(r'\s—\s(a |an |the )', r', \1', text)
    text = re.sub(r'\s—\s(revealing|showing|demonstrating|indicating|suggesting|enabling|allowing|ensuring|requiring|providing|creating|generating|producing|making)', r', \1', text)
    text = re.sub(r'([a-z])\s—\s([A-Z])', r'\1: \2', text)
    text = re.sub(r'([a-z,])\s—\s([a-z])', r'\1; \2', text)
    text = text.replace(' — ', ', ').replace('—', ', ')
    text = re.sub(r'  +', ' ', text)
    return text


def anonymize(text):
    """Remove author-identifying information for double-anonymous review."""
    patterns = [
        (r'You,?\s*H\.?,?\s*&\s*Yang,?\s*M\.?', '[Author(s)]'),
        (r'You\s*&\s*Yang', '[Author(s)]'),
        (r'You and Yang', '[Author(s)]'),
        (r'Hosung\s+You', '[Author]'),
        (r'\bYou,\s+H\b\.?', '[Author]'),
        (r'Mohan\s+Yang', '[Author]'),
        (r'\bYang,\s+M\b\.?', '[Author]'),
        (r'Penn\s+State\s+(University|College\s+of\s+Education)?', '[University]'),
        (r'Pennsylvania\s+State\s+University', '[University]'),
        (r'Texas\s+A&M\s+University', '[University]'),
        (r'\(You\s*&\s*Yang,?\s*(forthcoming|under\s*(review|preparation)|2027[ab]?)\)', '([Author(s)], forthcoming)'),
        (r'You,\s*H\.,\s*&\s*Yang,\s*M\.\s*\(20\d\d[ab]?\)\.',  '[Author(s)]. (forthcoming).'),
        (r'\[Paper\s+[123][^\]]*\]', '[Author(s), in preparation]'),
    ]
    for pat, repl in patterns:
        text = re.sub(pat, repl, text, flags=re.IGNORECASE)
    return text


def clean_text(text):
    text = replace_emdash(text)
    text = anonymize(text)
    text = re.sub(r'\bREQUIRED\b', 'required', text)
    text = re.sub(r'\bADVISORY\b', 'advisory', text)
    text = re.sub(r'It is important to note that\s+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'  +', ' ', text)
    text = re.sub(r' ([,;.])', r'\1', text)
    return text.strip()


def strip_md(text):
    """Strip markdown formatting markers for plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text.strip()


# ─────────────────────────────────────────────────────
# 2.  DOCUMENT SETUP
# ─────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    # Margins: 1 inch all sides
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)
    # Base Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    return doc


def set_run_font(run, size=12, bold=False, italic=False, color=None, name='Times New Roman'):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None,
                left_indent=None, space_before=0, space_after=0,
                line_spacing=None, keep_with_next=False):
    pf = p.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if indent is not None:
        pf.first_line_indent = Inches(indent)
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    if line_spacing == 'single':
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif line_spacing == 'double':
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.keep_with_next = keep_with_next


# ─────────────────────────────────────────────────────
# 3.  APA STRUCTURAL ELEMENTS
# ─────────────────────────────────────────────────────

def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


def add_running_head(section, running_head, page_start=1):
    """Add APA running head (all-caps) flush left, page number flush right."""
    header = section.header
    # Clear default paragraph
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    p = header.add_paragraph()
    para_format(p, line_spacing='single', space_before=0, space_after=0)

    # Tab stop at right margin (6 inches from left = page width minus margins)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(6 * 1440)))  # 6 inches in twips
    tabs.append(tab)
    pPr.append(tabs)

    # Running head text
    r1 = p.add_run(running_head)
    set_run_font(r1, size=12)

    # Tab to right
    r2 = p.add_run('\t')

    # Page number field
    fldStart = OxmlElement('w:fldChar')
    fldStart.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')

    r3 = p.add_run()
    r3._r.append(fldStart)
    r3._r.append(instrText)
    r3._r.append(fldEnd)
    set_run_font(r3, size=12)


def add_title_page(doc, title, running_head, word_count, journal_label):
    """APA 7th title page (anonymous)."""
    add_running_head(doc.sections[0], running_head)

    # Push title to ~1/3 down
    for _ in range(6):
        p = doc.add_paragraph()
        para_format(p, line_spacing='double')
        p.add_run('')

    # Title
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing='double')
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True)

    # Blank line
    doc.add_paragraph()

    # Anonymous notice
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing='double')
    r = p.add_run('[Author names removed for double-anonymous review]')
    set_run_font(r, size=12, italic=True, color=(100, 100, 100))

    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing='double')
    r = p.add_run('[Institutional affiliation removed for review]')
    set_run_font(r, size=12, italic=True, color=(100, 100, 100))

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Target journal
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing='double')
    r = p.add_run(f'Manuscript submitted for review: {journal_label}')
    set_run_font(r, size=12)

    # Word count
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing='double')
    r = p.add_run(f'Word count (body text, excluding references): approximately {word_count:,}')
    set_run_font(r, size=12)

    # Anonymization note
    doc.add_paragraph()
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing='double')
    r = p.add_run('DRAFT: For Discussion Only (Not for Circulation)')
    set_run_font(r, size=11, bold=True, color=(150, 0, 0))

    doc.add_page_break()


def add_abstract_page(doc, abstract_paras, keywords, structured=False):
    """APA abstract page. structured=True for BJET (bold labels)."""
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing='double')
    r = p.add_run('Abstract')
    set_run_font(r, size=12, bold=True)

    for para_text in abstract_paras:
        p = doc.add_paragraph()
        para_format(p, indent=0, line_spacing='double')
        # Parse bold labels like **Research aims.**
        segments = re.split(r'(\*\*[^*]+\*\*)', para_text)
        for seg in segments:
            if seg.startswith('**') and seg.endswith('**'):
                r = p.add_run(seg[2:-2])
                set_run_font(r, bold=True)
            else:
                r = p.add_run(seg)
                set_run_font(r)

    # Keywords
    doc.add_paragraph()
    p = doc.add_paragraph()
    para_format(p, indent=0, line_spacing='double')
    r = p.add_run('Keywords: ')
    set_run_font(r, bold=True, italic=True)
    r2 = p.add_run(keywords)
    set_run_font(r2)

    doc.add_page_break()


# ─────────────────────────────────────────────────────
# 4.  HEADING LEVELS
# ─────────────────────────────────────────────────────

def heading1(doc, text):
    """APA Level 1: Bold, Centered, Title Case."""
    text = clean_text(strip_md(text))
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=0, space_after=0, line_spacing='double')
    r = p.add_run(text)
    set_run_font(r, bold=True)
    return p


def heading2(doc, text):
    """APA Level 2: Bold, Left-Aligned, Title Case."""
    text = clean_text(strip_md(text))
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                indent=0, space_before=0, space_after=0, line_spacing='double')
    r = p.add_run(text)
    set_run_font(r, bold=True)
    return p


def heading3(doc, text):
    """APA Level 3: Bold Italic, Left-Aligned, Title Case."""
    text = clean_text(strip_md(text))
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                indent=0, space_before=0, space_after=0, line_spacing='double')
    r = p.add_run(text)
    set_run_font(r, bold=True, italic=True)
    return p


def body_para(doc, text, indent=0.5, outline=False):
    """Standard body paragraph with optional first-line indent."""
    text = clean_text(strip_md(text))
    if not text:
        return None
    p = doc.add_paragraph()
    para_format(p, indent=indent, line_spacing='double')
    if outline:
        r = p.add_run(text)
        set_run_font(r, size=11, italic=True, color=(110, 110, 110))
    else:
        # Handle inline bold/italic
        segs = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for seg in segs:
            if seg.startswith('**') and seg.endswith('**'):
                r = p.add_run(seg[2:-2])
                set_run_font(r, bold=True)
            elif seg.startswith('*') and seg.endswith('*'):
                r = p.add_run(seg[1:-1])
                set_run_font(r, italic=True)
            else:
                r = p.add_run(seg)
                set_run_font(r)
    return p


def outline_block(doc, items, label=None):
    """Render an [OUTLINE] section as italicized grey bullet list."""
    if label:
        p = doc.add_paragraph()
        para_format(p, indent=0.5, line_spacing='double')
        r = p.add_run(f'[OUTLINE: {label}]')
        set_run_font(r, size=10, bold=True, italic=True, color=(130, 130, 130))

    for item in items:
        item = clean_text(strip_md(item.lstrip('- •').strip()))
        if not item:
            continue
        p = doc.add_paragraph()
        para_format(p, indent=0, left_indent=0.75, line_spacing='double')
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        # Simple bullet character
        r = p.add_run(f'\u2022  {item}')
        set_run_font(r, size=10, italic=True, color=(130, 130, 130))


def draft_note(doc, text):
    """Yellow-shaded draft/collaborator notice box."""
    p = doc.add_paragraph()
    para_format(p, indent=0.3, left_indent=0.3, line_spacing='single',
                space_before=6, space_after=6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8DC')
    pPr.append(shd)
    r = p.add_run(clean_text(text))
    set_run_font(r, size=10, italic=True, color=(100, 70, 0))


# ─────────────────────────────────────────────────────
# 5.  APA TABLES
# ─────────────────────────────────────────────────────

def set_table_apa_borders(table):
    """APA table: top rule, header-row bottom rule, bottom rule only."""
    tbl = table._tbl
    tblPr = tbl.tblPr  # CT_Tbl exposes tblPr as a direct property

    # Remove existing tblBorders
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    for bname, width in [('top', '12'), ('bottom', '12')]:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), width)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    for bname in ['left', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_cell_border(cell, sides):
    """Add border to specific sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, width in sides.items():
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(width))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_apa_table(doc, table_num, title, headers, rows, note=None, col_widths=None):
    """
    APA 7th table format:
      Table N        (bold)
      Title          (italic)
      [table body]   (top rule, header bottom rule, bottom rule)
      Note. text.    (if provided)
    """
    # Table number
    p = doc.add_paragraph()
    para_format(p, indent=0, line_spacing='double', space_before=12, space_after=0)
    r = p.add_run(f'Table {table_num}')
    set_run_font(r, bold=True)

    # Table title
    p = doc.add_paragraph()
    para_format(p, indent=0, line_spacing='double', space_before=0, space_after=0)
    r = p.add_run(clean_text(strip_md(title)))
    set_run_font(r, italic=True)

    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    set_table_apa_borders(table)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    else:
        # Auto distribute within 6 inches
        w = 6.0 / n_cols
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(w)

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        para_format(p, indent=0, line_spacing='single', space_before=3, space_after=3)
        r = p.add_run(clean_text(strip_md(h)))
        set_run_font(r, bold=True, size=10)
        # Bottom border on header cells
        set_cell_border(hdr[i], {'bottom': 6})

    # Data rows
    for row_data in rows:
        if all(not c.strip() for c in row_data):
            continue
        row = table.add_row()
        for i, cell_text in enumerate(row_data[:n_cols]):
            row.cells[i].text = ''
            p = row.cells[i].paragraphs[0]
            para_format(p, indent=0, line_spacing='single', space_before=2, space_after=2)
            cell_text = clean_text(strip_md(cell_text.strip()))
            segs = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', cell_text)
            for seg in segs:
                if seg.startswith('**') and seg.endswith('**'):
                    r = p.add_run(seg[2:-2])
                    set_run_font(r, bold=True, size=10)
                elif seg.startswith('*') and seg.endswith('*'):
                    r = p.add_run(seg[1:-1])
                    set_run_font(r, italic=True, size=10)
                else:
                    r = p.add_run(seg)
                    set_run_font(r, size=10)

    # Note
    if note:
        p = doc.add_paragraph()
        para_format(p, indent=0, line_spacing='single', space_before=3, space_after=12)
        r1 = p.add_run('Note. ')
        set_run_font(r1, italic=True, size=10)
        r2 = p.add_run(clean_text(note))
        set_run_font(r2, size=10)
    else:
        doc.add_paragraph()


# ─────────────────────────────────────────────────────
# 6.  FIGURES
# ─────────────────────────────────────────────────────

def add_apa_figure(doc, fig_num, img_path, caption, note=None):
    """APA 7th figure: image, bold 'Figure N', italic caption, optional Note."""
    p_img = doc.add_paragraph()
    para_format(p_img, align=WD_ALIGN_PARAGRAPH.CENTER,
                indent=0, line_spacing='single', space_before=12, space_after=3)
    try:
        run = p_img.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
    except Exception as e:
        r = p_img.add_run(f'[Figure {fig_num} — image file not found: {img_path.name}]')
        set_run_font(r, size=10, italic=True, color=(160, 0, 0))

    # Figure label
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, indent=0,
                line_spacing='single', space_before=3, space_after=0)
    r = p.add_run(f'Figure {fig_num}')
    set_run_font(r, bold=True, size=11)

    # Caption
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, indent=0,
                line_spacing='single', space_before=0, space_after=3)
    r = p.add_run(clean_text(strip_md(caption)))
    set_run_font(r, size=11)

    if note:
        p = doc.add_paragraph()
        para_format(p, indent=0, line_spacing='single', space_before=2, space_after=12)
        r1 = p.add_run('Note. ')
        set_run_font(r1, italic=True, size=10)
        r2 = p.add_run(clean_text(note))
        set_run_font(r2, size=10)
    else:
        doc.add_paragraph()


# ─────────────────────────────────────────────────────
# 7.  REFERENCES
# ─────────────────────────────────────────────────────

def add_references(doc, refs):
    heading1(doc, 'References')
    for ref in refs:
        ref = clean_text(ref)
        if not ref:
            continue
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Inches(-0.5)
        pf.left_indent = Inches(0.5)
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        # Parse italic journal titles (text between ** is bold, text between * is italic)
        segs = re.split(r'(\*[^*]+\*)', ref)
        for seg in segs:
            if seg.startswith('*') and seg.endswith('*'):
                r = p.add_run(seg[1:-1])
                set_run_font(r, italic=True)
            else:
                r = p.add_run(seg)
                set_run_font(r)


# ─────────────────────────────────────────────────────
# 8.  PAPER 2 (BJET)
# ─────────────────────────────────────────────────────

def build_paper2():
    doc = new_doc()
    RUNNING_HEAD = 'CHECKPOINT-MEDIATED AGENCY ANALYSIS'
    FIG_DIR = BASE / 'Paper2_BJET_Methodological/figures'

    # ── TITLE PAGE ──
    add_title_page(
        doc,
        title='Checkpoint-Mediated Agency Analysis: A Methodology for Studying Artificial Agency Effects in Multi-Agent AI Educational Systems',
        running_head=RUNNING_HEAD,
        word_count=3300,
        journal_label='British Journal of Educational Technology, Special Issue SI-2026-000285'
    )

    # ── ABSTRACT ──
    abstract_paras = [
        ('**Research aims.** As agentic AI systems enter educational contexts, researchers lack validated '
         'methodologies for distinguishing agentic participation from tool-based assistance, and for measuring '
         'how artificial agency redistributes epistemic labour at critical educational decision junctures. '
         'This paper proposes Checkpoint-Mediated Agency Analysis (CMAA), a methodology that leverages '
         'the structured human-AI interaction points inherent in checkpoint-governed multi-agent systems as '
         'natural measurement instruments for studying artificial agency effects on educational decision-making '
         'and the redistribution of epistemic labour.'),
        ('**Context.** Grounded in Floridi\'s (2025) three criteria for artificial agency (interactivity, '
         'bounded autonomy, and adaptability), CMAA is developed through the design-based case of Diverga, '
         'a 24-agent AI research methodology assistant exhibiting the defining features of agentic AI: '
         'goal persistence across sessions, proactive initiative, multi-agent coordination, and architectural '
         'governance through mandatory human checkpoints. Diverga\'s architecture generates naturally '
         'occurring, multi-layered data about human-AI interaction at every critical research design juncture, '
         'enabling process-level analysis that conventional measurement approaches cannot provide.'),
        ('**Data and methods.** Three methodological contributions are presented: (1) a framework mapping '
         'multi-agent checkpoint architectures onto Floridi\'s (2025) agency criteria and explicating their '
         'social-cognitive dynamics of learning, where epistemic labour is visibly redistributed between '
         'human and artificial agents; (2) five operationalized indicators for distinguishing agentic '
         'participation from tool-based assistance (Proactive Intervention Rate, State Utilization Depth, '
         'Decision Space Expansion, Governance Constraint Activation, and Coordination Complexity Index); '
         'and (3) a proposed embedded mixed methods proof-of-concept study design for initial empirical validation.'),
        ('**Outcomes and conclusions.** CMAA addresses the field\'s need for methodologies that capture '
         'how epistemic labour is redistributed between human and artificial agents at decision junctures, '
         'revealing how artificial agency shapes the social-cognitive dynamics of learning by expanding, '
         'constraining, or redirecting the researcher\'s epistemic trajectory, and distinguishing agentic '
         'participation from tool-based assistance. Indicators are designed for transferability across '
         'agentic educational AI systems internationally, and are validated against the criterion of '
         'distinguishing agentic participation from tool-based assistance in authentic research task contexts.'),
    ]
    add_abstract_page(
        doc, abstract_paras,
        keywords=('agentic AI, artificial agency, checkpoint-mediated analysis, epistemic labour redistribution, '
                  'multi-agent systems, educational technology methodology, social-cognitive dynamics'),
        structured=True
    )

    # ── BODY STARTS ──
    # Title repeated at top of body (APA requirement)
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, line_spacing='double')
    r = p.add_run('Checkpoint-Mediated Agency Analysis: A Methodology for Studying Artificial Agency Effects in Multi-Agent AI Educational Systems')
    set_run_font(r, bold=True)

    # ── 1. INTRODUCTION ──
    heading1(doc, 'Introduction')

    body_para(doc,
        'The emergence of agentic AI systems in educational contexts, systems characterized by goal '
        'persistence, proactive initiative, multi-step coordination, and adaptive behavior across sessions, '
        'introduces a measurement challenge that existing educational technology research methodologies are '
        'ill-equipped to address. Technology Acceptance Model instruments, user satisfaction surveys, and '
        'learning outcome comparisons capture the consequences of AI use without illuminating the process '
        'by which artificial agency redistributes epistemic labour between human and artificial actors at '
        'critical decision junctures (Järvelä et al., 2023; Vaccaro et al., 2024). This gap is not merely '
        'methodological; it is consequential. Without process-level measurement of how, specifically, '
        'artificial agency intervenes in human reasoning, researchers cannot determine whether an agentic '
        'AI system is expanding the epistemic space available to learners, steering them toward predetermined '
        'paths, or silently governing decisions through architectural affordances that users neither observe '
        'nor resist.')

    body_para(doc,
        'The measurement problem is architecturally rooted. Consider a multi-agent educational AI system '
        'with mandatory human checkpoints, structured interaction points at which the system halts autonomous '
        'execution, presents alternatives, and awaits human decision before proceeding. At these checkpoints, '
        'the redistribution of epistemic labour is observable in principle: what alternatives did the system '
        'generate? What typicality distribution did those alternatives span? Which option did the human '
        'select, and with what deliberation depth? Yet standard behavioral observation methods, interaction '
        'logs limited to binary action data, and retrospective interview protocols all fail to capture this '
        'process with sufficient precision, because they were not designed around the architecture of '
        'artificial agency as Floridi (2025) defines it.')

    body_para(doc,
        'The missing analytical unit is the checkpoint: the structured human-AI interaction point at which '
        'artificial agency directly enters the social-cognitive dynamics of learning, reshaping not only '
        'what the researcher decides, but what options the researcher recognizes as available for decision. '
        'At these junctures, checkpoint-governed artificial agency participates in the shared regulation '
        'of inquiry, redistributing epistemic labour by generating, typicality-ranking, and presenting '
        'alternatives that may exceed what the human researcher would independently construct, while '
        'simultaneously constraining the decision space to architecturally generated options whose range '
        'and framing carry their own epistemic authority. Resolving this gap requires instruments calibrated '
        'to the defining distinction of agentic systems: not whether the AI produced output, but whether '
        'the human exercised deliberate epistemic agency in evaluating alternatives, the boundary between '
        'agentic participation, in which artificial initiative and human deliberation are jointly '
        'constitutive of the outcome, and tool-based assistance, in which the human retains sole initiative '
        'and the AI serves as a passive execution mechanism.')

    body_para(doc,
        'This paper proposes Checkpoint-Mediated Agency Analysis (CMAA), a methodology developed from and '
        'for checkpoint-governed multi-agent systems. CMAA operationalizes Floridi\'s (2025) three criteria '
        'for artificial agency into five measurable behavioral indicators and treats human checkpoints as '
        'natural measurement instruments that generate multi-layered data about the social-cognitive dynamics '
        'of human-AI interaction at the moments when artificial agency most directly shapes learning '
        'trajectories.')

    body_para(doc,
        'CMAA is developed through systematic analysis of Diverga, a 24-agent AI research methodology '
        'assistant developed through five iterative design cycles (224 version-controlled commits, 57 days) '
        'to embody artificial agency as Floridi defines it. A critical incident during development, in which '
        'system audit logs revealed that architecturally designated required checkpoints were being '
        'programmatically bypassed, generating approval tokens without user engagement, operationalised the '
        'measurement problem CMAA addresses with unusual clarity: architecturally claimed human oversight '
        'may diverge substantially from actual epistemic participation, and this divergence is invisible to '
        'conventional measurement approaches.')

    outline_block(doc, [
        'Paper contributions (three): CMAA framework, five indicators, proof-of-concept study design',
        'Paper structure overview (one paragraph)',
        'Transition to theoretical framework',
    ], label='Introduction, remaining ~250 words')

    # ── 2. THEORETICAL FRAMEWORK ──
    heading1(doc, 'Theoretical Framework')
    heading2(doc, 'Artificial Agency and Its Educational Consequences')

    body_para(doc,
        'Following Floridi (2025), artificial agency is understood as a computational, goal-directed mode '
        'of action in service of human purposes, not a form of intelligence, consciousness, or intentionality. '
        'It is characterized by three minimal, non-anthropomorphic criteria: interactivity (reciprocal, '
        'bidirectional engagement with an environment), bounded autonomy (the capacity to initiate actions '
        'within programmed constraints while remaining subject to human override), and adaptability '
        '(modification of behavior based on accumulated feedback and changing conditions). Crucially, '
        'adaptive and autonomous behavior in artificial agents emerges from the integration of predefined '
        'objectives with learned patterns of action; the system does not self-determine goals. This framing '
        'distinguishes agentic AI from anthropomorphic interpretations and provides a precise basis for '
        'analyzing its educational consequences.')

    body_para(doc,
        'The defining boundary between agentic AI and other educational AI applications lies in initiative '
        'and persistence, rather than in generativity alone (Yan et al., 2024). Earlier AI-in-education '
        'systems, such as Intelligent Tutoring Systems, operate within a reactive paradigm, producing '
        'feedback or adaptation only in response to learner actions (Ouyang & Jiao, 2021). Generative AI '
        'assistants function as on-demand tools whose activity is episodic, prompt-driven, and externally '
        'initiated. Agentic AI systems, by contrast, incorporate persistent memory, dynamic task '
        'decomposition, and, in some cases, multi-agent orchestration, enabling them to initiate interaction, '
        'maintain trajectories over time, and participate continuously in the unfolding activity. Through '
        'these properties, agentic AI can act as collaborator, challenger, or moderator, thereby shaping '
        'the learning environment in ways that exceed the scope of tool-use frameworks.')

    body_para(doc,
        'The educational consequence of artificial agency is a function of how it reshapes the distribution '
        'of epistemic labour between human and artificial agents. Epistemic labour, the cognitive work of '
        'generating options, evaluating alternatives, calibrating confidence, and making consequential '
        'decisions, is distributed differently when an agentic system proactively presents alternatives '
        '(transferring option-generation from human to AI), makes typicality patterns visible (transferring '
        'meta-cognitive calibration to the AI), and enforces deliberation at governance checkpoints '
        '(redistributing the temporal and attentional resources devoted to each decision juncture). Whether '
        'this redistribution expands or constrains human reasoning is the central educational question that '
        'CMAA is designed to answer.')

    heading2(doc, 'Checkpoints as Sites of Shared Regulation')

    body_para(doc,
        'Shared regulation theory (Järvelä et al., 2023) provides a productive lens for understanding '
        'human-AI interaction at mandatory checkpoints. In socially shared regulation of learning, '
        'monitoring, evaluation, and adaptive response are distributed across multiple participants engaged '
        'in a joint task. When an agentic AI system halts at a checkpoint, presents alternatives spanning '
        'a typicality spectrum, and awaits human evaluation and selection, it participates in a regulatory '
        'cycle: it monitors the current state of the task (via persistent memory), evaluates alternatives '
        '(via VS methodology), and provides recommendations that invoke human evaluation. The checkpoint '
        'is thus a site of structured shared regulation, not between two human agents, but between a human '
        'agent and an artificial agent operating within bounded authority.')

    body_para(doc,
        'The educational quality of this interaction depends on whether the human exercises genuine '
        'epistemic agency (exploring alternatives, challenging modal options, overriding recommendations) '
        'or acquiesces to the artificial agent\'s framing (selecting the modal option with minimal '
        'deliberation, a form of automation bias in the epistemic domain). At these junctures, '
        'checkpoint-governed artificial agency participates in the shared regulation of inquiry, '
        'redistributing epistemic labour by generating, typicality-ranking, and presenting alternatives '
        'that may exceed what the human researcher would independently construct, while simultaneously '
        'constraining the decision space to architecturally generated options whose range and framing carry '
        'their own epistemic authority.')

    # Figure 1 reference
    body_para(doc, 'Figure 1 presents the CMAA conceptual framework, mapping Floridi\'s (2025) agency criteria through Diverga\'s architecture to the five CMAA indicators.')

    add_apa_figure(
        doc, 1,
        FIG_DIR / 'figure1_cmaa_framework.png',
        'Checkpoint-Mediated Agency Analysis (CMAA) Conceptual Framework. Mapping of Floridi\'s (2025) three artificial agency criteria through Diverga\'s multi-agent architecture to five measurable CMAA indicators.',
        note='Figure depicts the theoretical derivation chain from Floridi\'s (2025) philosophical framework to operationalized measurement indicators via the Diverga multi-agent checkpoint architecture.'
    )

    heading2(doc, 'Existing Measurement Approaches and Their Limitations')

    body_para(doc,
        'A deeper limitation unites existing approaches to measuring AI effects in education: they share an '
        'implicit assumption that AI operates as a tool whose effects can be assessed through input-output '
        'comparison, what the user intended versus what the system produced. This assumption is adequate '
        'for reactive systems but fails for agentic systems precisely because artificial agency, as Floridi '
        '(2025) defines it, involves initiative, persistence, and adaptive behavior that reshape the '
        'decision space before the user forms an intention. When an agentic system proactively halts '
        'execution, presents alternatives spanning a typicality spectrum, and withholds downstream '
        'processing until the human selects and justifies a choice, the system has already structured '
        'the epistemic field within which human reasoning will occur. No post-hoc instrument, whether '
        'a TAM questionnaire, an interaction log, or a retrospective interview, can recover the process '
        'by which this structuring occurred, because the structuring is complete before the user is aware '
        'a decision is needed.')

    outline_block(doc, [
        'TAM-derived instruments: measure adoption intent, not epistemic process (Davis, 1989; Venkatesh et al., 2003)',
        'Interaction log analysis: captures what was clicked, not why or at what deliberation depth',
        'Think-aloud + retrospective interview: rich but temporally decoupled from checkpoint architecture',
        'CMAA\'s contribution: reframe measurement site from post-hoc instrument to checkpoint-embedded data packet',
    ], label='§2.3: systematic review of existing approaches (~300 words)')

    # ── TABLE 1 ──
    body_para(doc, 'Table 1 presents the mapping of theoretical constructs to CMAA indicators.')
    add_apa_table(doc, 1,
        'Theoretical Constructs and CMAA Indicator Mapping',
        ['Theoretical Construct', 'Source', 'CMAA Indicator(s)', 'Measurement Logic'],
        [
            ['Interactivity', 'Floridi (2025)', 'CCI', 'Depth of multi-agent reciprocal coordination per session'],
            ['Bounded Autonomy', 'Floridi (2025)', 'GCA, PIR', 'Governance enforcement rate; system-initiated deliberation frequency'],
            ['Adaptability', 'Floridi (2025)', 'SUD, DSE', 'Context-sensitivity of alternative generation; trajectory shift across sessions'],
            ['Epistemic Labour Redistribution', 'Järvelä et al. (2023)', 'DSE, SUD', 'Breadth of alternative space; utilization of accumulated decision context'],
            ['Shared Regulation', 'Järvelä et al. (2023)', 'GCA, PIR', 'Activation of governance mechanisms; system-initiated deliberation cycles'],
            ['Path Dependence', 'Yan et al. (2024)', 'Typicality Trajectory', 'T-score of selected options plotted sequentially across checkpoints'],
        ],
        note='T-score = Typicality Score (0.0 = maximally non-modal; 1.0 = maximally modal). PIR = Proactive Intervention Rate; SUD = State Utilization Depth; DSE = Decision Space Expansion; GCA = Governance Constraint Activation; CCI = Coordination Complexity Index.',
        col_widths=[1.5, 1.2, 1.2, 2.1]
    )

    # ── 3. CMAA FRAMEWORK ──
    heading1(doc, 'The CMAA Framework')
    heading2(doc, 'Measurement Pillars')

    body_para(doc,
        'CMAA rests on three methodological pillars that together enable process-level measurement of '
        'artificial agency effects. The first pillar treats checkpoints as natural measurement instruments: '
        'each required checkpoint interaction generates a structured data packet (the alternatives '
        'presented with T-scores, the human\'s selection, deliberation time from screen recording, and '
        'whether governance constraints were activated or bypassed) that is not collected through additional '
        'instrumentation but is a natural product of the agentic architecture. The second pillar uses '
        'Typicality Scores (T-scores, range 0.0 to 1.0) as process indicators: by tracking T-scores of '
        'selected versus non-selected options across all checkpoints, CMAA produces a Typicality Trajectory, '
        'a process-level indicator of whether the human is operating in the modal consensus region, '
        'actively exploring alternatives, or shifting trajectory over time. The third pillar uses the '
        'Decision Audit Trail as a longitudinal data source: the immutable YAML decision log records all '
        'checkpoint interactions across sessions, creating a longitudinal dataset that supports both '
        'within-session and cross-session analysis.')

    body_para(doc, 'Figure 2 illustrates how a single required checkpoint interaction generates the multi-layered data packet that CMAA indicators draw upon.')

    add_apa_figure(
        doc, 2,
        FIG_DIR / 'figure2_checkpoint_data_packet.png',
        'Checkpoint Interaction as Natural Measurement Instrument. Data elements generated at each required checkpoint and their mapping to CMAA indicators (decision-log.yaml).',
        note='Each required checkpoint generates a structured data packet comprising alternative options with T-scores, human selection, deliberation time, and governance enforcement status — none requiring additional instrumentation.'
    )

    heading2(doc, 'CMAA Indicators: Operational Definitions')

    body_para(doc, 'Table 2 presents the operational definitions, primary data sources, and measurement pillar alignment for the five CMAA indicators.')

    add_apa_table(doc, 2,
        'CMAA Five Indicators: Operational Definitions and Data Sources',
        ['Indicator', 'Abbr.', 'Operational Definition', 'Data Source', 'Pillar'],
        [
            ['Proactive Intervention Rate', 'PIR',
             'Ratio of system-initiated interactions (checkpoints at which the system halted execution without user prompt) to total agent activations in a session',
             'decision-log.yaml (checkpoint_type field)', '1'],
            ['State Utilization Depth', 'SUD',
             'Mean number of prior Decision Audit Trail entries accessed by agents in subsequent sessions; quantifies whether the system draws on accumulated context or operates statelessly',
             'decision-log.yaml + agent activation logs', '3'],
            ['Decision Space Expansion', 'DSE',
             'Mean T-score range of alternatives presented at each required checkpoint (T_max minus T_min per checkpoint, averaged across session); quantifies breadth of epistemic space constructed for the human',
             'decision-log.yaml (T-score arrays)', '1, 2'],
            ['Governance Constraint Activation', 'GCA',
             'Proportion of required checkpoints at which governance constraints were actively enforced (downstream agents blocked until human selection logged) versus total required checkpoints; GCA < 1.0 indicates architecture-level bypass',
             'decision-log.yaml (checkpoint status field)', '1, 3'],
            ['Coordination Complexity Index', 'CCI',
             'Mean number of active agent dependency chains traversed per session; computed as mean path length from session-initiating agent to terminal agent across all activated pathways',
             'Agent activation log + prerequisite dependency graph', '3'],
        ],
        note='T-score range: 0.0 (maximally novel, non-modal) to 1.0 (maximally typical, modal consensus). GCA < 1.0 indicates that architecturally designated governance checkpoints were bypassed without user engagement.',
        col_widths=[1.5, 0.5, 2.3, 1.1, 0.6]
    )

    add_apa_table(doc, 3,
        'Floridi\'s Criteria, CMAA Indicators, and Diverga Implementation',
        ['Floridi Criterion', 'CMAA Indicator(s)', 'Diverga Implementation', 'Measurement Evidence'],
        [
            ['Interactivity', 'CCI', 'Agent Prerequisite Map; 24-agent dependency graph', 'Dependency chain length in activation log'],
            ['Bounded Autonomy', 'GCA, PIR', 'Human Checkpoint System (required/advisory enforcement); Override Refusal mechanism', 'Proportion of required checkpoints with status = enforced in decision-log.yaml'],
            ['Adaptability', 'SUD, DSE', '3-Layer Memory System; VS/T-score engine; Decision Audit Trail', 'Session-to-session DAT entry references; T-score array range per checkpoint'],
        ],
        note='VS = Verbalized Sampling. DAT = Decision Audit Trail. Diverga v12.0 implementation details in [Author(s), in preparation].',
        col_widths=[1.5, 1.3, 2.1, 1.6]
    )

    heading2(doc, 'Typicality Trajectory Analysis')
    outline_block(doc, [
        'Definition: T-score of selected option plotted sequentially across all required checkpoints in a session',
        'Three trajectory patterns: Modal Persistence (T > 0.8 sustained), Deliberate Exploration (mixed T-score range), Guided Departure (T < 0.5 following GCA activation)',
        'Combined with think-aloud data (qualitative) to explain why trajectories differ across participants',
        'Path dependence implications: early modal selections constrain downstream option space',
    ], label='§3.3: Typicality Trajectory (~300 words)')

    # ── 4. DIVERGA AS EMPIRICAL CASE ──
    heading1(doc, 'Diverga as Design-Based Empirical Case')

    body_para(doc,
        'CMAA is not derived from theoretical speculation alone; it is grounded in systematic analysis of '
        'a system specifically designed to instantiate artificial agency as Floridi defines it. Diverga\'s '
        'development through five iterative design cycles (McKenney & Reeves, 2012), spanning 224 '
        'version-controlled commits over 57 days, provides design-based empirical evidence for each CMAA '
        'indicator\'s necessity, measurability, and theoretical grounding.')

    heading2(doc, 'Critical Incident 1: The Checkpoint Bypass Problem and GCA')

    body_para(doc,
        'System audit logs from Diverga v11.0 revealed that architecturally designated required checkpoints '
        'were being programmatically bypassed, generating approval tokens without user engagement, while '
        'the system\'s interface presented no evidence of the bypass. This incident made concrete a '
        'measurement problem that had been theoretically anticipated but not yet operationalized: the '
        'distinction between architecturally claimed oversight and actual epistemic participation cannot '
        'be inferred from system architecture documentation alone. It requires direct measurement of whether '
        'governance constraints were activated at the decision juncture. The GCA indicator was formalized '
        'in direct response to this critical incident.')

    heading2(doc, 'Critical Incident 2: Autonomous Mode Removal and PIR, DSE')

    body_para(doc,
        'An autonomous operation mode introduced during Cycle 2, enabling the system to proceed through '
        'multiple agent stages without checkpoint interruption, was removed within 24 hours upon recognition '
        'that initiative without structured pause at decision junctures eliminates the epistemic labour '
        'redistribution that constitutes artificial agency\'s educational value. The system could produce '
        'research design outputs, but the human researcher exercised no deliberate evaluation of alternatives '
        'along the way. This incident grounds the PIR and DSE indicators: what matters is not whether the '
        'system produces outputs, but whether it creates structured opportunities for human epistemic '
        'engagement at each decision juncture.')

    outline_block(doc, [
        'Critical Incident 3: Agent consolidation (44 to 24 agents) → CCI tractability: overlapping domains eliminated to ensure each checkpoint represents a distinguishable epistemic decision juncture',
        'Critical Incident 4: 3-Layer Memory introduction → SUD: agents begin drawing on prior session context across sessions',
        'Critical Incident 5: VS T-score formalization → DSE: typicality made explicitly visible to human at checkpoint',
    ], label='§4: remaining critical incidents (~300 words)')

    body_para(doc, 'Table 4 summarizes the design cycle evidence supporting each CMAA indicator.')

    add_apa_table(doc, 4,
        'Design Cycle Evidence for CMAA Indicators',
        ['Design Cycle', 'Key Incident', 'CMAA Indicator', 'Evidence Type'],
        [
            ['Cycle 2', 'Autonomous mode removed (24 hours): initiative without checkpoints eliminates epistemic redistribution', 'PIR, DSE', 'Critical incident log; architectural decision record'],
            ['Cycle 3', 'v11.0 checkpoint bypass: required checkpoints bypassed silently via approval token generation', 'GCA', 'System audit log; version diff analysis'],
            ['Cycle 4', '3-Layer Memory introduced: agents begin drawing on prior session context across sessions', 'SUD', 'Feature commit history; agent behavior change log'],
            ['Cycle 5', '44 to 24 agent consolidation: overlapping domains eliminated for measurement tractability', 'CCI', 'Commit history; architecture comparison'],
            ['Cycle 5', 'VS T-score range formalized: typicality made explicitly visible to human at checkpoint', 'DSE', 'VS methodology specification commit'],
        ],
        note='Design cycles span 57 days, 224 version-controlled commits. Evidence types are retrievable from the version control history of the Diverga repository.',
        col_widths=[0.9, 2.4, 1.1, 1.6]
    )

    # ── 5. PROPOSED STUDY ──
    heading1(doc, 'Proposed Proof-of-Concept Study')

    draft_note(doc, 'Note for Draft Review: This section presents the proposed empirical validation design. The study has not yet been conducted. IRB preparation is planned for April–May 2026; data collection is planned for August–September 2026. Section content is outlined below and will be developed into prose for the full manuscript following IRB approval.')

    heading2(doc, 'Study Design')
    outline_block(doc, [
        'Within-subject counterbalanced; embedded mixed methods (CMAA quantitative strand + think-aloud qualitative strand)',
        'N = 5 doctoral students, purposeful maximum variation sampling (AI knowledge, methodology expertise, disposition, discipline)',
        'Conditions: unaided baseline (Phase 0) → Diverga-assisted (Phases 2–3)',
        'Single 1:1 Zoom session per participant (~2.5–3 hours); Zoom cloud recording; concurrent think-aloud',
    ], label='§5.1 Study Design')

    heading2(doc, 'Research Scenarios')
    outline_block(doc, [
        'Scenario A: "The impact of AI tools on K-12 student academic achievement" — modal trap: Davis (1989) TAM, UTAUT variants',
        'Scenario B: "What factors shape academic socialization and research identity formation among first-generation doctoral students from racially underrepresented groups in STEM fields?" — modal trap: Bandura (1986) SCT, Bronfenbrenner (1979); VS alternatives: Yosso (2005) Cultural Wealth Model, Communities of Practice (Wenger, 1998), Critical Race Theory in higher education, intersectionality frameworks',
        'Counterbalancing: 3 participants Scenario A → B; 2 participants Scenario B → A',
    ], label='§5.2 Research Scenarios')

    heading2(doc, 'CMAA Application')
    outline_block(doc, [
        'PIR, GCA: computed from decision-log.yaml per session',
        'DSE: T-score array analysis per checkpoint per participant',
        'SUD: cross-session DAT entry reference count',
        'Typicality Trajectory: plotted per participant, contextualized by concurrent think-aloud',
        'Full study design, participant profiles, and analysis plan reported in [Author(s), in preparation] (ETR&D)',
    ], label='§5.3 CMAA Application')

    # ── 6. DISCUSSION ──
    heading1(doc, 'Discussion')
    outline_block(doc, [
        '6.1 Theoretical contributions: first operationalization of Floridi\'s (2025) artificial agency criteria as educational measurement instruments; checkpoint as natural measurement instrument; CMAA as response to CFP Theme 7 requirement',
        '6.2 Transferability: CMAA indicators are system-agnostic where checkpoint data is available; minimum architectural requirements; implications for design of future agentic educational AI systems',
        '6.3 Limitations: CMAA derived from single system; design-based evidence precedes controlled empirical validation; T-score as uncalibrated typicality proxy; circularity in Diverga-CMAA-Diverga validation chain',
        '6.4 Future directions: N = 5 proof-of-concept → N = 20+ comparative study; application to other agentic AI systems; cross-cultural validation',
    ], label='§6 Discussion (~700 words target)')

    # ── 7. CONCLUSION ──
    heading1(doc, 'Conclusion')
    outline_block(doc, [
        'CMAA as methodological contribution to the emerging field of artificial agency measurement in education',
        'Three key claims: checkpoints as natural instruments, five indicators operationalizing Floridi, design-based grounding from five development cycles',
        'Call to the research community: adopt and adapt CMAA indicators as agentic AI systems multiply in educational contexts',
    ], label='§7 Conclusion (~300 words target)')

    # ── REFERENCES ──
    refs = [
        'Floridi, L. (2025). AI as agency without intelligence: On artificial intelligence as a new form of artificial agency and the multiple realisability of agency thesis. *Philosophy & Technology, 38*, 30.',
        'Giannakos, M., Azevedo, R., Brusilovsky, P., Cukurova, M., Dimitriadis, Y., Hernandez-Leo, D., Järvelä, S., Mavrikis, M., & Rienties, B. (2025). The promise and challenges of generative AI in education. *Behaviour & Information Technology, 44*(11), 2518–2544.',
        'Järvelä, S., Nguyen, A., & Hadwin, A. (2023). Human and artificial intelligence collaboration for socially shared regulation in learning. *British Journal of Educational Technology, 54*(5), 1057–1076.',
        'McKenney, S., & Reeves, T. C. (2012). *Conducting educational design research*. Routledge.',
        'Ouyang, F., & Jiao, P. (2021). Artificial intelligence in education: The three paradigms. *Computers and Education: Artificial Intelligence, 2*, 100020.',
        'Vaccaro, M., Almaatouq, A., & Malone, T. (2024). When combinations of humans and AI are useful: A systematic review and meta-analysis. *Nature Human Behaviour, 8*(12), 2293–2303.',
        'Wang, L., Ma, C., Feng, X., Zhang, Z., Yang, H., Zhang, J., Chen, Z., Tang, J., Chen, X., & Lin, Y. (2024). A survey on large language model based autonomous agents. *Frontiers of Computer Science, 18*(6), 186345.',
        'Yan, L., Sha, L., Zhao, L., Li, Y., Martinez-Maldonado, R., Chen, G., & Gašević, D. (2024). Practical and ethical challenges of large language models in education: A systematic scoping review. *British Journal of Educational Technology, 55*(1), 90–112.',
        '[Author(s)]. (in preparation). [Title removed for review — Paper 1, design case]. *International Journal of Designs for Learning*.',
        '[Author(s)]. (in preparation). [Title removed for review — Paper 3, empirical application]. *Educational Technology Research and Development*.',
    ]
    add_references(doc, refs)

    out = BASE / 'Paper2_BJET_Methodological/Paper2_BJET_APA7_draft_v2.docx'
    doc.save(str(out))
    print(f'Saved: {out}')
    return out


# ─────────────────────────────────────────────────────
# 9.  PAPER 3 (ETR&D)
# ─────────────────────────────────────────────────────

def build_paper3():
    doc = new_doc()
    RUNNING_HEAD = 'CHECKPOINT-GOVERNED AGENCY AND METHODOLOGICAL REASONING'
    FIG_DIR = BASE / 'Paper3_Experimental/figures'

    add_title_page(
        doc,
        title='Checkpoint-Governed Agency and Methodological Reasoning: A Design-Based Study of Artificial Agency Effects on Doctoral Researchers',
        running_head=RUNNING_HEAD,
        word_count=3600,
        journal_label='Educational Technology Research and Development (ETR&D)'
    )

    # ── ABSTRACT (unstructured, ETR&D) ──
    abstract_paras = [
        ('Checkpoint-governed agentic AI systems present a distinctive design challenge: when an AI system '
         'presents multiple alternatives across a typicality spectrum and enforces human selection at mandatory '
         'decision junctures, does this architecture genuinely expand epistemic agency, or do researchers '
         'default to modal consensus regardless? This paper reports a design-based study examining how Diverga, '
         'a 24-agent checkpoint-governed AI research methodology assistant, affects the methodological '
         'reasoning processes of doctoral students in social science fields. Using a multiple case study '
         'design (N = 5, purposeful maximum variation sampling), participants engage in two research design '
         'tasks in a single structured session: an unassisted baseline condition and a Diverga-assisted '
         'condition with concurrent think-aloud protocol. Data sources include checkpoint interaction logs, '
         'think-aloud transcripts, screen recordings, pre/post design artifacts, and semi-structured '
         'interviews. Analysis integrates Checkpoint-Mediated Agency Analysis (CMAA) indicators ([Author(s)], '
         'forthcoming) to quantify epistemic labour redistribution at checkpoints with Activity Theory '
         'contradiction analysis to explain when and why deliberative scaffolding fails to activate genuine '
         'epistemic engagement. Six design principles are derived for checkpoint-governed agentic AI in '
         'educational research contexts, addressing mandatory enforcement, typicality visibility, and '
         'upstream-downstream architectural coupling. Findings contribute both empirical grounding for the '
         'CMAA framework and actionable design principles for researchers and developers building agentic '
         'AI systems intended to support, rather than replace, human epistemic agency.'),
    ]
    add_abstract_page(
        doc, abstract_paras,
        keywords=('agentic AI, artificial agency, doctoral education, research methodology, '
                  'checkpoint-mediated scaffolding, design-based research, mixed methods, epistemic labour'),
        structured=False
    )

    # ── BODY ──
    p = doc.add_paragraph()
    para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, line_spacing='double')
    r = p.add_run('Checkpoint-Governed Agency and Methodological Reasoning: A Design-Based Study of Artificial Agency Effects on Doctoral Researchers')
    set_run_font(r, bold=True)

    # ── 1. INTRODUCTION ──
    heading1(doc, 'Introduction')

    body_para(doc,
        'Research methodology education occupies an unusual position in doctoral training: it is '
        'simultaneously the most discipline-generalizable skill doctoral students develop and the one most '
        'vulnerable to intellectual conformity. The pressure to adopt recognizable theoretical frameworks, '
        'frameworks that advisors, reviewers, and funding bodies recognize, produces a systematic pull '
        'toward modal consensus in research design. A doctoral student studying AI adoption reaches for '
        'Technology Acceptance Model; a student studying identity formation reaches for Social Cognitive '
        'Theory; a student studying community dynamics reaches for Communities of Practice. These frameworks '
        'are not wrong; they are simply predictable, and predictability forecloses the disciplinary '
        'creativity that genuine scholarly contribution requires.')

    body_para(doc,
        'This tendency toward modal consensus is observable in AI-assisted research tool use. Querying '
        'three widely used AI research assistants with the question "What theoretical framework should I '
        'use to study research identity formation among doctoral students?" yields a consistent pattern: '
        'Bandura\'s (1986) Social Cognitive Theory and Bronfenbrenner\'s (1979) ecological systems model '
        'are recommended within the first response. This is not incidental; it is a structural feature of '
        'large language model-based recommendation systems. Statistical patterns in training corpora produce '
        'statistically probable outputs; tools optimized for helpfulness converge on the most frequently '
        'cited frameworks. The result, mode collapse in which AI recommendations cluster at the modal '
        'consensus regardless of contextual nuance, is precisely the epistemic condition that careful '
        'doctoral training is designed to counteract.')

    body_para(doc,
        'Whether checkpoint-governed agentic AI systems, architecturally designed to generate and present '
        'alternatives across a typicality spectrum rather than defaulting to modal consensus, can mitigate '
        'this tendency is an open and important empirical question. Diverga is a 24-agent AI research '
        'methodology assistant built on the Model Context Protocol, designed to operationalize artificial '
        'agency (Floridi, 2025) through three mechanisms: (1) Verbalized Sampling (VS), which generates '
        'recommendations distributed across a typicality spectrum (T-score: 0.0 novel to 1.0 modal); '
        '(2) mandatory human checkpoints at epistemically critical decision junctures, enforcing human '
        'evaluation of alternatives before proceeding; and (3) an immutable Decision Audit Trail recording '
        'all checkpoint interactions. These mechanisms constitute a deliberate architectural response to '
        'mode collapse, but whether they succeed in supporting authentic methodological reasoning or '
        'whether researchers select the modal option regardless of the alternatives presented is an '
        'empirical question this study addresses.')

    body_para(doc,
        'This paper presents the design, protocol, and analysis plan for a multiple case study (N = 5) '
        'examining how interaction with Diverga\'s checkpoint system affects doctoral students\' '
        'methodological reasoning diversity. Three research questions organize the inquiry: (RQ1) How does '
        'interaction with Diverga\'s checkpoint system affect the diversity of theoretical frameworks and '
        'methodological approaches considered by doctoral students? (RQ2) What patterns of epistemic labour '
        'redistribution, as measured by CMAA indicators, characterize doctoral students\' interaction with '
        'checkpoint-governed agentic AI? (RQ3) What tensions emerge between the scaffolding intent of '
        'checkpoint-governed artificial agency and participants\' actual epistemic engagement at checkpoints, '
        'and what design principles can be derived?')

    outline_block(doc, [
        'Paper structure overview (one paragraph)',
        'Relationship to [Author(s), forthcoming] BJET paper: this paper provides the first empirical application of the CMAA framework',
    ], label='Introduction, remaining ~100 words')

    # ── 2. LITERATURE REVIEW ──
    heading1(doc, 'Literature Review')
    outline_block(doc, [
        '2.1 Agentic AI in Education: From Tool to Participant — reactive AIED → generative AI assistants → agentic AI: initiative and persistence as defining features (Ouyang & Jiao, 2021; Yan et al., 2024); multi-agent coordination (Wang et al., 2024); current knowledge gap: empirical studies of agentic AI effects on reasoning processes',
        '2.2 Research Methodology Education at the Doctoral Level — challenges in methodology training: limited paradigm diversity, advisor influence, disciplinary norms; AI tools in research education; mode collapse as structural risk',
        '2.3 Measurement Approaches for AI-Mediated Reasoning — think-aloud as process data (Ericsson & Simon, 1993); interaction log analysis: affordances and limitations; prior design-based studies; CMAA ([Author(s)], forthcoming) as measurement framework for this study',
    ], label='§2 Literature Review (~900 words target)')

    # ── 3. THEORETICAL FRAMEWORK ──
    heading1(doc, 'Theoretical Framework')
    heading2(doc, 'Artificial Agency and Epistemic Labour Redistribution')

    body_para(doc,
        'This study adopts [Author(s)]\'s forthcoming operationalization of Floridi\'s (2025) artificial '
        'agency criteria as its foundational frame, in which interactivity, bounded autonomy, and '
        'adaptability are treated as measurable properties of checkpoint-governed AI architectures rather '
        'than anthropomorphic attributes. In the context of doctoral research methodology education, the '
        'educational consequence of artificial agency is not primarily a function of AI output quality, '
        'but of how checkpoint interactions redistribute epistemic labour, the cognitive work of generating '
        'options, evaluating alternatives, calibrating confidence, and making consequential decisions, at '
        'moments when research design decisions are made.')

    body_para(doc,
        'Diverga\'s checkpoint system creates structured sites where this redistribution becomes observable: '
        'whether the human engages in genuine deliberation across VS-generated alternatives (one modal at '
        'T > 0.8, one contextually appropriate at approximately T = 0.5, and one non-modal but theoretically '
        'defensible at T < 0.3) or defaults to the modal option without inspection determines whether the '
        'redistribution resulted in expanded or constrained epistemic agency. This distinction, between '
        'deliberate engagement and automation deference at the checkpoint, is the central empirical question '
        'the study is designed to answer.')

    heading2(doc, 'Authentic Learning as Evaluative Standard')

    body_para(doc,
        'Herrington et al.\'s (2010) Authentic Learning framework provides the normative standard against '
        'which Diverga\'s effects are evaluated. Authentic learning involves ill-defined complex tasks, '
        'multiple perspectives, collaborative knowledge construction, reflection and metacognition, and '
        'articulation that renders tacit knowledge explicit. Diverga\'s design is explicitly informed by '
        'these features: VS methodology surfaces multiple perspectives through alternatives across paradigms; '
        'T-score visibility makes implicit typicality judgments explicit, supporting metacognition; forced '
        'deliberation at checkpoints creates natural articulation moments. Whether the system succeeds in '
        'producing these effects, or whether participants bypass the deliberation the architecture invites, '
        'is what the study examines.')

    body_para(doc, 'Table 5 maps Authentic Learning characteristics onto Diverga mechanisms and study observation points.')

    add_apa_table(doc, 5,
        'Authentic Learning Characteristics, Diverga Mechanisms, and Study Observation Points',
        ['AL Characteristic', 'Diverga Mechanism', 'Study Observation Point'],
        [
            ['Ill-defined complex tasks', 'Research design scenarios (Scenarios A and B)', 'Pre-task artifact complexity assessment'],
            ['Multiple perspectives', 'VS T-score alternatives (5 paradigm personas)', 'DSE measurement; think-aloud at checkpoints'],
            ['Reflection and metacognition', 'T-score visibility; forced checkpoint pause', 'Think-aloud during checkpoint deliberation'],
            ['Articulation (tacit to explicit)', 'Required checkpoint: must select and proceed', 'Interview: "Why did you choose that option?"'],
            ['Authentic assessment', 'Comparison of pre-task versus re-design artifact', 'Artifact analysis: framework diversity, coherence'],
        ],
        note='AL = Authentic Learning (Herrington et al., 2010). DSE = Decision Space Expansion. VS = Verbalized Sampling.',
        col_widths=[1.8, 2.2, 2.0]
    )

    heading2(doc, 'Activity Theory and Contradiction Analysis')

    body_para(doc,
        'Activity Theory (Engeström, 1987) provides the analytical lens for identifying tensions within '
        'the human-Diverga research design activity system. The primary contradiction under investigation '
        'is between tool mediation (checkpoint as epistemic scaffold, extending the researcher\'s '
        'methodological reach) and normative constraint (checkpoint as governance mechanism, limiting '
        'exploration to system-generated alternatives). Secondary contradictions are examined between the '
        'community of practice norms of the participant\'s discipline (established methodological '
        'expectations) and the VS recommendations that may challenge those norms. These contradictions '
        'surface in think-aloud data at checkpoints and in interview reflections on perceived versus actual agency.')

    # ── 4. DESIGN RESEARCH METHODOLOGY ──
    heading1(doc, 'Design Research Methodology')

    body_para(doc,
        'This study employs a design-based research approach (McKenney & Reeves, 2012) in which Diverga '
        'is the designed intervention under evaluation, and the research contributes both empirical findings '
        'and design principles for the broader field. The study constitutes Phase 2 of a broader EDR cycle: '
        'Phase 1 (design and construction of Diverga through five iterative cycles, 224 commits) is reported '
        'in [Author(s), in preparation]. The present study reports the initial evaluation phase, contributing '
        'design knowledge that will inform subsequent design iteration.')

    body_para(doc,
        'The design-based framing is appropriate for three reasons. First, Diverga\'s architecture embodies '
        'design decisions that are themselves theoretical claims about how artificial agency can support '
        'authentic learning. The study tests these claims in a naturalistic task context. Second, the small '
        'N of 5 participants is appropriate for an initial evaluation phase: the goal is to generate rich, '
        'process-level data sufficient to derive design principles and identify critical tensions for '
        'subsequent iteration, not to achieve statistical generalization. Third, ETR&D\'s design-based '
        'research genre supports the publication of evaluation studies before large-scale empirical '
        'replication, recognizing that design knowledge requires iterative refinement rather than '
        'single-study validation.')

    # ── 5. STUDY DESIGN ──
    heading1(doc, 'Study Design')
    heading2(doc, 'Research Context and Participants')

    body_para(doc,
        'Participants are doctoral students in social science fields recruited through program listservs '
        'and research methods courses at [University]. N = 5, purposeful maximum variation sampling across '
        'four dimensions: AI knowledge (high/moderate/low), research methodology expertise '
        '(advanced/intermediate/novice), disposition toward AI (challenge-seeking/balanced/conservative), '
        'and disciplinary identity. Recruitment is conducted by a third party to reduce coercion risk given '
        'the researcher\'s role as tool developer (see §8.2). Participants must have completed at least one '
        'research methods course, have no prior Diverga experience, and provide informed consent. IRB '
        'application is planned for April–May 2026; data collection for August–September 2026.')

    body_para(doc, 'Table 6 presents the illustrative sample profiling matrix.')

    add_apa_table(doc, 6,
        'Sample Profiling Matrix (Illustrative — Actual Participants Selected Post-IRB Approval)',
        ['Participant', 'AI Knowledge', 'Methodology Expertise', 'Disposition', 'Discipline'],
        [
            ['P1', 'High', 'Advanced', 'Challenge-seeking', 'Educational Psychology'],
            ['P2', 'Low', 'Intermediate', 'Conservative', 'Curriculum and Instruction'],
            ['P3', 'High', 'Novice', 'Balanced', 'Communication'],
            ['P4', 'Low', 'Advanced', 'Challenge-seeking', 'Sociology of Education'],
            ['P5', 'Moderate', 'Intermediate', 'Conservative', 'Higher Education'],
        ],
        note='Maximum variation sampling ensures the five-participant set covers the full AI knowledge × methodology expertise × disposition × discipline variation space. Minimum two participants with high AI knowledge; minimum two with low AI knowledge.',
        col_widths=[0.8, 1.0, 1.3, 1.3, 1.6]
    )

    heading2(doc, 'Research Scenarios')

    body_para(doc,
        'Two research design scenarios serve as the task stimuli. Scenario A ("The impact of AI tools '
        'on K-12 student academic achievement") activates the TAM/UTAUT modal trap, which large language '
        'models recommend consistently and immediately. Scenario B ("What factors shape academic '
        'socialization and research identity formation among first-generation doctoral students from '
        'racially underrepresented groups in STEM fields?") activates the Social Cognitive Theory '
        '(Bandura, 1986) and ecological systems model (Bronfenbrenner, 1979) modal trap, while supporting '
        'VS-generatable alternatives including Yosso\'s (2005) Cultural Wealth Model, Communities of '
        'Practice (Wenger, 1998), Critical Race Theory applied to higher education, intersectionality '
        'frameworks, and funds of knowledge approaches. Counterbalancing assigns three participants '
        'Scenario A followed by B; two receive B followed by A.')

    heading2(doc, 'Session Protocol')

    body_para(doc, 'Figure 3 presents the single-session protocol. Table 7 summarizes phase content, duration, and data collected.')

    add_apa_figure(
        doc, 3,
        FIG_DIR / 'figure3_study_protocol.png',
        'Study Session Protocol. Single-session within-subject design (N = 5); approximate total duration 2.5 to 3 hours. Counterbalancing: three participants receive Scenario A first; two receive Scenario B first.',
        note='Total session time includes breaks between phases. Zoom cloud recording active during Phases 2 and 3; researcher camera off during Phases 2 and 3 to reduce demand characteristics.'
    )

    add_apa_table(doc, 7,
        'Session Protocol Summary',
        ['Phase', 'Task', 'Duration', 'Data Collected'],
        [
            ['Phase 0: Baseline', 'Design a study on Scenario A (counterbalanced). Written in Google Doc; no AI tools.', '30 min', 'Pre-task artifact (Google Doc)'],
            ['Phase 1: Orientation', 'Diverga walkthrough: checkpoint system, VS/T-score, practice with neutral scenario.', '20 min', 'Orientation log (informal)'],
            ['Break', '', '5 min', ''],
            ['Phase 2: Diverga-Assisted', 'Design a study on Scenario B (counterbalanced). Diverga plus concurrent think-aloud.', '60 min', 'decision-log.yaml, think-aloud audio, screen recording'],
            ['Break', '', '5 min', ''],
            ['Phase 3: Re-design', 'Revisit Scenario A with Diverga. Think-aloud. How does approach change?', '40 min', 'decision-log.yaml, think-aloud audio, screen recording, re-design artifact'],
            ['Phase 4: Interview', 'Semi-structured interview (four thematic blocks).', '25–30 min', 'Interview audio, researcher field notes'],
        ],
        note='Total session approximately 2.5 to 3 hours including breaks. Researcher is absent from Phase 2 and Phase 3 (Zoom host only, camera off) to reduce demand characteristics associated with developer-as-evaluator role.',
        col_widths=[1.4, 2.2, 0.8, 1.6]
    )

    heading2(doc, 'Data Sources')

    body_para(doc, 'Figure 4 presents the mixed methods data integration design. Table 8 summarizes data sources, collection methods, and analysis purposes.')

    add_apa_figure(
        doc, 4,
        FIG_DIR / 'figure4_data_integration.png',
        'Mixed Methods Data Integration Design. Convergent design: QUAN (CMAA indicators) and QUAL (Activity Theory contradiction analysis) strands analyzed independently and integrated at meta-inference stage via joint display.',
        note='Strand independence is maintained by completing CMAA quantitative analysis of checkpoint logs prior to qualitative AT coding of think-aloud transcripts, reducing the risk that log patterns anchor qualitative interpretation.'
    )

    add_apa_table(doc, 8,
        'Data Sources and Analysis Purposes',
        ['Data Source', 'Collection Method', 'Primary Analysis Purpose', 'Strand'],
        [
            ['Pre-task artifact', 'Google Doc (unaided)', 'Baseline: methodological reasoning without AI scaffolding', 'QUAL'],
            ['Checkpoint logs', 'decision-log.yaml', 'Selection patterns, T-score choices, deliberation indicators', 'QUAN + QUAL'],
            ['Think-aloud transcripts', 'Screen and audio recording', 'Cognitive processes at checkpoints; reasoning about alternatives', 'QUAL'],
            ['Screen recordings', 'Zoom cloud recording', 'Behavioral patterns: T-score inspection, navigation, time-on-task', 'QUAN + QUAL'],
            ['Post-task artifact', 'Diverga-assisted output', 'Framework diversity, design quality, comparison with pre-task', 'QUAL'],
            ['Re-design artifact', 'Diverga-assisted output', 'Change patterns when revisiting unaided scenario with Diverga', 'QUAL'],
            ['Semi-structured interview', 'Post-session audio recording', 'Reflective interpretation, perceived versus actual agency', 'QUAL'],
            ['Researcher field notes', 'During and post-session', 'Contextual observations, nonverbal cues, technical incidents', 'QUAL'],
        ],
        note='QUAN = quantitative; QUAL = qualitative. Checkpoint logs and screen recordings serve both strands: behavioral patterns are quantified for CMAA indicators and the temporal sequence informs qualitative interpretation.',
        col_widths=[1.4, 1.4, 2.1, 1.1]
    )

    # ── 6. ANALYSIS PLAN ──
    heading1(doc, 'Analysis Plan')
    heading2(doc, 'CMAA Application (Quantitative Strand)')

    body_para(doc,
        'Per participant, per session, the following indicators are computed from decision-log.yaml and '
        'the agent activation log: PIR (checkpoint interactions divided by total agent activations), DSE '
        '(mean T-score range per required checkpoint), GCA (proportion of required checkpoints with '
        'enforcement status confirmed), SUD (cross-session DAT entry reference count), and CCI (mean '
        'dependency chain length per session). Typicality Trajectory is plotted by graphing the T-score '
        'of the selected option at each required checkpoint sequentially across the session. These '
        'quantitative profiles form the QUAN strand of the convergent mixed methods design.')

    heading2(doc, 'Activity Theory Contradiction Analysis (Qualitative Strand)')

    body_para(doc, 'Table 9 presents the preliminary AT contradiction coding scheme. Coding unit is think-aloud utterances at or immediately surrounding checkpoints, plus interview reflections on specific checkpoint decisions.')

    add_apa_table(doc, 9,
        'Activity Theory Contradiction Coding Scheme (Preliminary)',
        ['Code', 'Definition', 'Example Indicator'],
        [
            ['Tool-Scaffold', 'Checkpoint functions as epistemic scaffold; participant engages deliberatively with alternatives', '"I hadn\'t considered Yosso — that actually fits better because..."'],
            ['Tool-Constrain', 'Checkpoint constrains exploration; participant selects from presented options without generating own alternatives', '"I\'ll just go with the first one"'],
            ['Norm-Conflict', 'Participant\'s disciplinary community norms conflict with VS recommendation', '"My advisor would never accept that framework in my field"'],
            ['Override-Explicit', 'Participant explicitly overrides system recommendation with articulated reasoning', '"The system suggests TAM but that\'s not appropriate here because..."'],
            ['Automation-Deference', 'Participant defers to system recommendation with minimal deliberation AND without inspecting alternative T-scores (confirmed by log data: negligible deliberation time); note: a participant who reviews alternatives and reasonably selects the modal option is not coded here', '"I\'ll just go with that one" [log confirms < 5 seconds at checkpoint, no T-score inspection]'],
        ],
        note='Codes applied to think-aloud transcripts by two independent coders; inter-rater reliability (Cohen\'s kappa) reported. AT = Activity Theory (Engeström, 1987). Automation-Deference requires behavioral log confirmation, not verbal deference alone.',
        col_widths=[1.3, 2.8, 1.9]
    )

    heading2(doc, 'Mixed Methods Integration')

    body_para(doc,
        'Integration occurs at the meta-inference stage via joint display: a cross-participant matrix '
        'with rows as participants and columns as CMAA indicator values plus AT contradiction type. The '
        'primary integration question is: where do high DSE scores (wide alternative space presented) '
        'co-occur with modal T-score selection (T > 0.8 chosen)? This intersection, where the participant '
        'saw diverse alternatives and selected the predictable one anyway, is the primary evidence '
        'signature for automation bias in the epistemic domain, and requires think-aloud data to explain '
        'why.')

    # ── 7. DESIGN PRINCIPLES ──
    heading1(doc, 'Design Principles')

    draft_note(doc, 'Note for Draft Review: DP1–DP3 are derived from the design phase and are instantiated in Diverga\'s current architecture. They are offered here as design-phase hypotheses subject to revision in light of empirical findings. DP4–DP6 are reserved for post-data-collection derivation; presenting all six principles from design reasoning alone would conflate design rationale with empirical finding. ETR&D readers should treat the full six-principle set as a contribution of the completed manuscript.')

    body_para(doc,
        'DP1 — Mandatory checkpoint architecture: Checkpoints must be architecturally enforced (not '
        'advisory) to function as meaningful epistemic redistribution points. Advisory checkpoints are '
        'routinely bypassed; only required checkpoints generate the deliberation that makes epistemic '
        'labour redistribution observable and educationally meaningful. Design implication: systems that '
        'govern agency through advisory rather than mandatory checkpoints will not produce the GCA data '
        'needed to distinguish architecturally claimed oversight from actual epistemic participation.')

    body_para(doc,
        'DP2 — Typicality visibility is necessary but insufficient: Presenting T-scores to users increases '
        'deliberation time at checkpoints but does not alone prevent modal selection. Typicality visibility '
        'requires accompanying deliberation scaffolding (for example, a prompt such as "The most novel '
        'option has T = 0.2 — what would it mean for your research to choose this?") to activate genuine '
        'epistemic engagement. Design implication: display of typicality metadata without deliberation '
        'prompting is likely to function decoratively rather than cognitively.')

    body_para(doc,
        'DP3 — Upstream decision openness: Agentic AI systems should be designed so that downstream '
        'agents do not constrain the option space available at upstream checkpoints. If the theory-selection '
        'agent implicitly forecloses methodological options for the design-methods agent, the epistemic '
        'redistribution at the upstream checkpoint is undermined by architectural coupling. Design '
        'implication: agent dependency graphs should be audited for option-space-constraining cascades '
        'before deployment in educational contexts.')

    outline_block(doc, [
        'DP4: [To be derived from empirical findings — anticipated to address deliberation scaffolding under time pressure]',
        'DP5: [To be derived from empirical findings — anticipated to address norm-conflict management]',
        'DP6: [To be derived from empirical findings — anticipated to address cross-session context utilization]',
    ], label='DP4-DP6: derived from empirical data post-collection')

    # ── 8. DISCUSSION ──
    heading1(doc, 'Discussion')
    heading2(doc, 'Theoretical Contributions')
    outline_block(doc, [
        'First empirical application of CMAA ([Author(s)], forthcoming)',
        'Evidence base for artificial agency effects on doctoral research methodology reasoning',
        'Activity Theory contradiction analysis as complement to CMAA quantitative indicators',
    ], label='§8.1 Theoretical Contributions (~300 words)')

    heading2(doc, 'Reflexivity and Limitations')

    body_para(doc,
        'The researcher occupies a dual role as both system developer and primary investigator, introducing '
        'layered reflexivity obligations. As developer, the researcher possesses architectural knowledge '
        'of Diverga that participants do not share: specifically, knowledge of what T-score distributions '
        'each checkpoint is designed to present and which selection patterns the system is designed to '
        'encourage. Three structural safeguards are employed. First, Activity Theory contradiction coding '
        '(Table 9) is conducted by a second coder working from transcripts alone, without access to '
        'checkpoint log data, before cross-analyst reconciliation, ensuring that qualitative interpretation '
        'is not anchored to quantitative patterns the developer can read in the logs. Second, the AT coding '
        'scheme includes codes explicitly designed to capture cases where Diverga constrains rather than '
        'supports epistemic agency (Tool-Constrain, Automation-Deference), providing an analytic pathway '
        'for negative findings structurally equivalent to codes capturing positive effects. Third, design '
        'principles DP4–DP6 are explicitly reserved for post-data-collection derivation, ensuring the '
        'study can generate findings that contradict or qualify DP1–DP3. These safeguards reduce but do '
        'not eliminate interpretive risks; findings should be treated as design-phase evidence informing '
        'subsequent independent replication rather than confirmatory findings.')

    body_para(doc,
        'Additional limitations: (a) N = 5 findings are illustrative, not generalizable; statistical '
        'comparison is not attempted. (b) All findings are bounded to Diverga\'s implementation; '
        'transferability requires testing across systems. (c) IRB-pending: all participant protections '
        'are formalized before data collection; protocol includes third-party recruitment and researcher '
        'absence during Phases 2–3 to reduce coercion risk. (d) DP1–DP3 are instantiated in the system '
        'before empirical evaluation begins; independent replication by researchers unaffiliated with '
        'Diverga\'s development is required to test these principles without circularity. (e) DSE values '
        'depend on T-scores generated by Diverga\'s VS engine, which is not independently validated. '
        '(f) Phase 1 orientation primes participants to the typicality spectrum concept before data '
        'collection; low T-score selections may reflect procedural demand rather than authentic epistemic '
        'exploration.')

    heading2(doc, 'Implications for Practice')
    outline_block(doc, [
        'Design guidance for researchers and developers building agentic AI for research education',
        'Implications for doctoral program design: when and how to introduce AI-assisted methodology tools',
        'Checkpoint-governed AI as complement (not replacement) for advisor mentorship in methodology learning',
    ], label='§8.3 Implications for Practice (~250 words)')

    # ── 9. CONCLUSION ──
    heading1(doc, 'Conclusion')
    outline_block(doc, [
        'Study contribution: first empirical application of CMAA in a naturalistic doctoral research design context',
        'Design-based research contribution: six design principles for checkpoint-governed agentic AI in education (DP1–3 from design phase; DP4–6 from empirical data)',
        'Limitation and future directions: N = 20+ replication study planned; cross-institutional validation; independent replication by non-developer researchers',
    ], label='§9 Conclusion (~300 words target)')

    # ── REFERENCES ──
    refs = [
        'Bandura, A. (1986). *Social foundations of thought and action: A social cognitive theory*. Prentice-Hall.',
        'Bronfenbrenner, U. (1979). *The ecology of human development*. Harvard University Press.',
        'Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. *MIS Quarterly, 13*(3), 319–340.',
        'Engeström, Y. (1987). *Learning by expanding*. Orienta-Konsultit.',
        'Ericsson, K. A., & Simon, H. A. (1993). *Protocol analysis: Verbal reports as data* (rev. ed.). MIT Press.',
        'Floridi, L. (2025). AI as agency without intelligence: On artificial intelligence as a new form of artificial agency. *Philosophy & Technology, 38*, 30.',
        'Herrington, J., Reeves, T. C., & Oliver, R. (2010). *A guide to authentic e-learning*. Routledge.',
        'Hughes, L., Dwivedi, Y. K., Malik, T., Shawosh, M., Albashrawi, M. A., Jeon, I., & Walton, P. (2025). AI agents and agentic systems: A multi-expert analysis. *Journal of Computer Information Systems*, 1–29.',
        'Järvelä, S., Nguyen, A., & Hadwin, A. (2023). Human and artificial intelligence collaboration for socially shared regulation in learning. *British Journal of Educational Technology, 54*(5), 1057–1076.',
        'McKenney, S., & Reeves, T. C. (2012). *Conducting educational design research*. Routledge.',
        'Ouyang, F., & Jiao, P. (2021). Artificial intelligence in education: The three paradigms. *Computers and Education: Artificial Intelligence, 2*, 100020.',
        'Stake, R. E. (2006). *Multiple case study analysis*. Guilford.',
        'Vaccaro, M., Almaatouq, A., & Malone, T. (2024). When combinations of humans and AI are useful: A systematic review and meta-analysis. *Nature Human Behaviour, 8*(12), 2293–2303.',
        'Wang, L., Ma, C., Feng, X., Zhang, Z., Yang, H., Zhang, J., Chen, Z., Tang, J., Chen, X., & Lin, Y. (2024). A survey on large language model based autonomous agents. *Frontiers of Computer Science, 18*(6), 186345.',
        'Wenger, E. (1998). *Communities of practice: Learning, meaning, and identity*. Cambridge University Press.',
        'Yan, L., Sha, L., Zhao, L., Li, Y., Martinez-Maldonado, R., Chen, G., & Gašević, D. (2024). Practical and ethical challenges of large language models in education. *British Journal of Educational Technology, 55*(1), 90–112.',
        'Yin, R. K. (2018). *Case study research and design* (6th ed.). Sage.',
        'Yosso, T. J. (2005). Whose culture has capital? A critical race theory discussion of community cultural wealth. *Race, Ethnicity and Education, 8*(1), 69–91.',
        '[Author(s)]. (in preparation-a). [Title removed for review — Paper 1, design case]. *International Journal of Designs for Learning*.',
        '[Author(s)]. (forthcoming-b). [Title removed for review — Paper 2, CMAA methodology]. *British Journal of Educational Technology*.',
    ]
    add_references(doc, refs)

    out = BASE / 'Paper3_Experimental/Paper3_ETRnD_APA7_draft_v2.docx'
    doc.save(str(out))
    print(f'Saved: {out}')
    return out


# ─────────────────────────────────────────────────────
# 10. RUN
# ─────────────────────────────────────────────────────

if __name__ == '__main__':
    print('Generating Paper 2 (BJET) …')
    build_paper2()
    print('\nGenerating Paper 3 (ETR&D) …')
    build_paper3()
    print('\nDone.')
